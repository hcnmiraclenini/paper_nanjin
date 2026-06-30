#!/usr/bin/env python3
"""Synchronize Traffic boundary wording into Word deliverables."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
PAPER_DOCX = ROOT / "paper/2026-0268-基于多专家融合的铁路客流多尺度预测方法.docx"
PAPER_COPY = ROOT / "paper/论文终稿.docx"
RESPONSE_DOCX = ROOT / "审稿意见逐条回复稿.docx"

PAPER_OLD = (
    "为评估模型在不同交通时序任务中的结构鲁棒性，在Traffic公开数据集（美国旧金山湾区高速公路小时级交通流量）上开展独立验证。"
)
PAPER_NEW = (
    "为评估模型在不同交通时序任务中的结构鲁棒性，本文保留Traffic公开数据集（美国旧金山湾区高速公路小时级交通流量）的既有补充基准记录。"
)
PAPER_EXTRA = (
    "需要说明的是，当前仓库未包含可端到端复现实验的Traffic原始数据文件，因此本节数值只作为现有Traffic补充基准记录，"
    "不作为当前仓库从原始Traffic数据重新训练复现的证据。"
)

RESPONSE_AUTO_AUDIT = (
    "自动审计补充：新增 scripts/audit_traffic_boundary.py，扫描工作区是否存在可复现实验用的 Traffic 原始数据文件，"
    "并检查论文、Word稿和审稿回复是否写入“现有补充基准记录、非主结论证据、重构小时级窗口、MAE非最优、不能直接外推”等边界。"
    "审计报告为 docs/experiments/Traffic跨粒度边界自动审计_20260701.md。该审计通过后，修订稿不再把 Traffic 写成当前仓库端到端重跑的独立实验。"
)
RESPONSE_TRAFFIC_RECHECK = (
    "Traffic边界复核：修订稿进一步将 2.8 节首句由“开展独立验证”改为“保留既有补充基准记录”，"
    "并明确当前仓库未包含可端到端复现实验的 Traffic 原始数据文件。自动审计已确认文稿保留“Traffic不作为主结论证据、"
    "MAE并非最优、不能直接外推稳定跨域泛化”的边界。"
)


def set_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "STSong")
    if run.font.size is None:
        run.font.size = Pt(10.5)


def normalize_fonts(doc: Document) -> None:
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                set_font(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            set_font(run)


def insert_after(paragraph: Paragraph, text: str) -> None:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = paragraph.style
    run = new_para.add_run(text)
    set_font(run)


def replace_or_append_extra(doc: Document) -> None:
    if any("当前仓库未包含可端到端复现实验的Traffic原始数据文件" in para.text for para in doc.paragraphs):
        return
    for para in doc.paragraphs:
        if PAPER_OLD in para.text:
            para.text = para.text.replace(PAPER_OLD, PAPER_NEW)
            if PAPER_EXTRA not in para.text:
                para.text = para.text + PAPER_EXTRA
            for run in para.runs:
                set_font(run)
            return
    for para in doc.paragraphs:
        if PAPER_NEW in para.text:
            if PAPER_EXTRA not in para.text:
                para.text = para.text + PAPER_EXTRA
                for run in para.runs:
                    set_font(run)
            return
    for para in doc.paragraphs:
        if "为评估结构鲁棒性，Traffic公开数据集实验采用domain adaptation framing" in para.text:
            para.text = para.text + PAPER_EXTRA
            for run in para.runs:
                set_font(run)
            return
    raise RuntimeError("Traffic paragraph not found in paper docx")


def insert_after_anchor(doc: Document, anchor: str, text: str, marker: str) -> None:
    if any(marker in para.text for para in doc.paragraphs):
        return
    for para in doc.paragraphs:
        if anchor in para.text:
            insert_after(para, text)
            return
    raise RuntimeError(f"Anchor not found in response docx: {anchor}")


def main() -> int:
    paper = Document(PAPER_DOCX)
    replace_or_append_extra(paper)
    normalize_fonts(paper)
    paper.save(PAPER_DOCX)
    shutil.copy2(PAPER_DOCX, PAPER_COPY)

    response = Document(RESPONSE_DOCX)
    insert_after_anchor(
        response,
        "该边界已写入",
        RESPONSE_AUTO_AUDIT,
        "自动审计补充",
    )
    insert_after_anchor(
        response,
        "论文同时删除过强跨域表述，将 Traffic 结果限定为结构鲁棒性初步验证。",
        RESPONSE_TRAFFIC_RECHECK,
        "Traffic边界复核",
    )
    normalize_fonts(response)
    response.save(RESPONSE_DOCX)
    print("updated Traffic boundary wording in Word deliverables")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
