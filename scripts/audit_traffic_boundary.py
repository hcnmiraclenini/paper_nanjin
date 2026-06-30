#!/usr/bin/env python3
"""Audit Traffic cross-resolution evidence boundaries.

The current repository contains legacy Traffic benchmark numbers but no raw
Traffic data artifact that would allow an end-to-end rerun. This script makes
that boundary explicit and checks the manuscript/reply texts do not turn the
legacy record into an over-claimed cross-domain generalization result.
"""

from __future__ import annotations

import json
import os
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent
ARTIFACT_JSON = ROOT / "docs/experiments/artifacts/traffic_boundary_audit_20260701.json"
REPORT_MD = ROOT / "docs/experiments/Traffic跨粒度边界自动审计_20260701.md"

RAW_SUFFIXES = {".csv", ".h5", ".hdf", ".hdf5", ".npz", ".npy", ".pkl", ".parquet"}
SKIP_DIRS = {".git", "__pycache__", ".pytest_cache", "node_modules", "checkpoints"}
TEXT_FILES = {
    "readme": ROOT / "README.md",
    "paper_md": ROOT / "paper/论文终稿.md",
    "response_md": ROOT / "审稿意见逐条回复稿.md",
    "expert_response_md": ROOT / "专家意见逐条修改说明.md",
    "traffic_boundary_md": ROOT / "docs/experiments/Traffic跨粒度证据边界_20260701.md",
    "final_consistency_md": ROOT / "docs/experiments/最终一致性核验_20260701.md",
}
DOCX_FILES = {
    "paper_docx": ROOT / "paper/2026-0268-基于多专家融合的铁路客流多尺度预测方法.docx",
    "response_docx": ROOT / "审稿意见逐条回复稿.docx",
}

REQUIRED_BOUNDARY_GROUPS = [
    ("legacy_record", ["现有Traffic补充基准记录", "现有 Traffic 补充基准记录", "既有四步长平均记录", "现有 Traffic 表格数值", "既有四预测步长平均结果记录"]),
    ("not_main_evidence", ["不作为论文主结论证据", "不作为主结论证据", "主结论仍以南京南站真实日粒度客流", "论文主结论仍以", "主结论和最终三指标优势只依赖", "主结论和三指标优势只以"]),
    ("reconstructed_features", ["重构小时级窗口", "重构时间窗口", "小时级输入窗口"]),
    ("mae_boundary", ["MAE指标上不优于", "MAE不是最优", "MAE并非最优", "MAE不优于"]),
    ("no_direct_transfer", ["不能直接外推", "不直接复用", "不主张铁路日粒度模型可直接", "不声称铁路日粒度模型可直接", "不能证明日粒度铁路模型可直接"]),
]

BANNED_NON_NEGATED = [
    "跨域泛化能力较强",
    "稳定跨域泛化能力",
    "直接跨分辨率迁移",
    "所有指标全面优于",
    "从原始数据重新训练复现",
    "开展独立验证",
]
NEGATION_MARKERS = ["不", "不能", "不得", "不宜", "不主张", "尚不足以", "并非", "仅", "只", "改为", "避免", "删除"]


def portable(path: Path) -> str:
    try:
        return str(path.resolve().relative_to(ROOT))
    except ValueError:
        try:
            return str(path.resolve().relative_to(WORKSPACE))
        except ValueError:
            return str(path.resolve())


def read_docx(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def find_raw_traffic_candidates(root: Path) -> list[str]:
    candidates: list[str] = []
    for current, dirs, files in os.walk(root):
        dirs[:] = [d for d in dirs if d not in SKIP_DIRS]
        current_path = Path(current)
        # Avoid spending time in generated caches and paper output folders.
        if any(part in {"docs", "paper", "logs", "results"} for part in current_path.parts):
            dirs[:] = []
            continue
        for file_name in files:
            path = current_path / file_name
            lower = file_name.lower()
            if "traffic" in lower and path.suffix.lower() in RAW_SUFFIXES:
                candidates.append(portable(path))
    return sorted(candidates)


def missing_groups(text: str, groups: list[tuple[str, list[str]]]) -> list[str]:
    missing = []
    for label, choices in groups:
        if not any(choice in text for choice in choices):
            missing.append(label)
    return missing


def non_negated_hits(text: str) -> list[dict]:
    hits = []
    for line_no, line in enumerate(text.splitlines(), start=1):
        for phrase in BANNED_NON_NEGATED:
            if phrase not in line:
                continue
            if any(marker in line for marker in NEGATION_MARKERS):
                continue
            hits.append({"line": line_no, "phrase": phrase, "text": line.strip()})
    return hits


def main() -> int:
    texts = {name: path.read_text(encoding="utf-8") for name, path in TEXT_FILES.items()}
    texts.update({name: read_docx(path) for name, path in DOCX_FILES.items()})

    raw_candidates = find_raw_traffic_candidates(WORKSPACE)
    no_raw_data = len(raw_candidates) == 0

    required_docs = ["paper_md", "response_md", "expert_response_md", "traffic_boundary_md", "final_consistency_md", "paper_docx", "response_docx"]
    boundary_missing = {
        name: missing_groups(texts[name], REQUIRED_BOUNDARY_GROUPS)
        for name in required_docs
    }
    boundary_missing = {name: miss for name, miss in boundary_missing.items() if miss}

    overclaims = {
        name: non_negated_hits(text)
        for name, text in texts.items()
    }
    overclaims = {name: hits for name, hits in overclaims.items() if hits}

    checks = [
        {
            "name": "no_raw_traffic_data_artifact_found",
            "passed": no_raw_data,
            "detail": raw_candidates,
        },
        {
            "name": "traffic_boundary_written_in_core_texts",
            "passed": not boundary_missing,
            "detail": boundary_missing,
        },
        {
            "name": "traffic_overclaim_phrase_scan",
            "passed": not overclaims,
            "detail": overclaims,
        },
    ]
    passed = all(check["passed"] for check in checks)
    report = {
        "passed": passed,
        "protocol": (
            "Read-only audit. It scans the workspace for raw Traffic data artifacts and checks manuscript "
            "wording boundaries. It performs no model training, tuning, or test-set evaluation."
        ),
        "workspace": portable(WORKSPACE),
        "raw_suffixes": sorted(RAW_SUFFIXES),
        "raw_traffic_candidates": raw_candidates,
        "checks": checks,
    }
    ARTIFACT_JSON.parent.mkdir(parents=True, exist_ok=True)
    ARTIFACT_JSON.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        "# Traffic 跨粒度边界自动审计（2026-07-01）",
        "",
        f"总体结果：{'通过' if passed else '未通过'}",
        "",
        "## 审计口径",
        "",
        "- 本审计只扫描文件和文稿，不训练模型、不调参、不评估测试集。",
        "- 搜索范围为项目工作区，候选 Traffic 原始数据文件限定为名称包含 `traffic` 且扩展名为 csv/h5/hdf/hdf5/npz/npy/pkl/parquet 的文件。",
        "- 若未发现原始数据，则论文只能将 Traffic 写成既有补充基准记录或结构鲁棒性边界，不能写成当前仓库可端到端复现的独立实验。",
        "",
        "## 检查结果",
        "",
        "| 检查项 | 结果 | 摘要 |",
        "| --- | --- | --- |",
    ]
    for check in checks:
        detail = check["detail"]
        if isinstance(detail, list):
            summary = f"{len(detail)} item(s)"
        elif isinstance(detail, dict):
            summary = f"{len(detail)} file(s) with hits"
        else:
            summary = str(detail)
        lines.append(f"| {check['name']} | {'通过' if check['passed'] else '失败'} | {summary} |")

    lines.extend(
        [
            "",
            "## 结论",
            "",
            "- 当前仓库仍不能把 Traffic 写成从原始数据完整重跑的实验；该边界已由自动审计固定。",
            "- 论文可保留 Traffic 作为跨时间粒度结构鲁棒性的补充记录，但主结论、全部指标优势和无泄露实验均以南京南站真实双向日粒度客流为准。",
            "- 若后续补齐 Traffic 原始数据，应重新记录数据哈希、预处理、验证集选模和最终一次测试结果，再更新本审计。",
        ]
    )
    REPORT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")

    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if passed else 1


if __name__ == "__main__":
    raise SystemExit(main())
