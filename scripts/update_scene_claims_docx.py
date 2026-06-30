#!/usr/bin/env python3
"""Insert scene-stratified audit wording into final Word deliverables."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
PAPER_DOCX = ROOT / "paper/2026-0268-基于多专家融合的铁路客流多尺度预测方法.docx"
PAPER_COPY = ROOT / "paper/论文终稿.docx"
RESPONSE_DOCX = ROOT / "审稿意见逐条回复稿.docx"

PAPER_SCENE_PARAGRAPH = (
    "进一步地，本文按可预知日历场景进行现代强基线分层审计。该审计同样只读取固定测试预测，"
    "工作日/周末标签由真实测试日期生成，不参与训练、调参、checkpoint选择或尺度校准。"
    "当前测试窗口包含152个工作日和60个周末样本；Strict RAMR-VE在工作日场景下的MAPE/MSE/MAE为"
    "18.1153%/0.038406/0.135967，在周末场景下为19.6382%/0.029723/0.142826，"
    "两类场景的三项指标均优于DLinear、PatchTST、iTransformer、TimeMixer和FreEformer。"
    "该结果从预测误差层面补充说明，日历场景特征和场景条件化门控并非只改善总体均值，"
    "而是在工作日和周末两种可检验场景中均保持优势；但当前测试窗口无节假日样本，"
    "因此本文不外推节假日场景结论。"
)

PAPER_CONCLUSION_OLD = (
    "配对预测审计中三指标同时更优的日期级bootstrap概率最低为89.40%。"
    "STL分解、门控权重统计和消融实验进一步支持"
)
PAPER_CONCLUSION_NEW = (
    "配对预测审计中三指标同时更优的日期级bootstrap概率最低为89.40%。"
    "工作日/周末场景分层审计进一步显示，Strict RAMR-VE在两类可检验日历场景下的三项指标均优于现代强基线。"
    "STL分解、门控权重统计和消融实验进一步支持"
)

RESPONSE_SCENE_GATE = (
    "场景误差补充：进一步新增工作日/周末场景分层现代强基线审计。该审计只读取固定测试预测，"
    "工作日/周末标签由真实测试日期生成，不训练、不调参、不重新选模。测试窗口包含152个工作日和60个周末；"
    "Strict RAMR-VE在工作日上的MAPE/MSE/MAE为18.1153%/0.038406/0.135967，在周末上为"
    "19.6382%/0.029723/0.142826，两类场景三项指标均优于DLinear、PatchTST、iTransformer、"
    "TimeMixer和FreEformer。由于测试窗口无节假日样本，修订稿不外推节假日结论。"
)

RESPONSE_SCENE_METRIC = (
    "场景分层审计：针对“日历场景输入是否真正有效”的疑问，新增工作日/周末现代强基线分层对比。"
    "Strict RAMR-VE在152个工作日样本上的MAPE/MSE/MAE为18.1153%/0.038406/0.135967，"
    "在60个周末样本上为19.6382%/0.029723/0.142826；两类场景下三项指标均优于五个现代强基线。"
    "该结果与门控统计中“工作日短期专家权重高、周末分布偏移专家权重高”的机制解释相互印证，"
    "但不作为测试集调参依据。"
)


def set_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "STSong")
    if run.font.size is None:
        run.font.size = Pt(10.5)


def normalize_fonts(doc: Document) -> None:
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                set_font(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            set_font(run)


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = paragraph.style
    run = new_para.add_run(text)
    set_font(run)
    return new_para


def insert_after_first(doc: Document, anchor: str, text: str, duplicate_marker: str) -> bool:
    if any(duplicate_marker in para.text for para in doc.paragraphs):
        return False
    for para in doc.paragraphs:
        if anchor in para.text:
            insert_after(para, text)
            return True
    raise RuntimeError(f"Anchor not found: {anchor}")


def replace_in_paragraphs(doc: Document, old: str, new: str) -> None:
    if any(new in para.text for para in doc.paragraphs):
        return
    for para in doc.paragraphs:
        if old in para.text:
            para.text = para.text.replace(old, new)
            for run in para.runs:
                set_font(run)
            return
    raise RuntimeError(f"Text to replace not found: {old}")


def update_paper() -> None:
    doc = Document(PAPER_DOCX)
    insert_after_first(
        doc,
        "因此本文将该结果表述为“点估计和日期级配对秩检验均支持优势”",
        PAPER_SCENE_PARAGRAPH,
        "18.1153%/0.038406/0.135967",
    )
    replace_in_paragraphs(doc, PAPER_CONCLUSION_OLD, PAPER_CONCLUSION_NEW)
    normalize_fonts(doc)
    doc.save(PAPER_DOCX)
    shutil.copy2(PAPER_DOCX, PAPER_COPY)


def update_response() -> None:
    doc = Document(RESPONSE_DOCX)
    insert_after_first(
        doc,
        "该结果说明门控并非简单平均融合，而是将高频相邻日变化和峰值突发强度分配给不同专家。",
        RESPONSE_SCENE_GATE,
        "场景误差补充",
    )
    insert_after_first(
        doc,
        "高日际变化日上 Strict RAMR-VE 的 MAPE 为 24.8108%",
        RESPONSE_SCENE_METRIC,
        "场景分层审计",
    )
    normalize_fonts(doc)
    doc.save(RESPONSE_DOCX)


def main() -> int:
    update_paper()
    update_response()
    print("updated Word deliverables with scene-stratified audit claims")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
