#!/usr/bin/env python3
"""Audit reviewer-comment completion coverage.

This script checks that all 15 reviewer comments have explicit evidence files
and wording coverage in the paper/reply package. It is read-only with respect
to model outputs and performs no training, tuning, checkpoint selection, or
test-set evaluation.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
ARTIFACT_JSON = ROOT / "docs/experiments/artifacts/reviewer_completion_audit_20260701.json"
REPORT_MD = ROOT / "docs/experiments/审稿意见完成度自动审计_20260701.md"

TEXT_FILES = {
    "readme": ROOT / "README.md",
    "paper_md": ROOT / "paper/论文终稿.md",
    "response_md": ROOT / "审稿意见逐条回复稿.md",
    "expert_response_md": ROOT / "专家意见逐条修改说明.md",
    "completion_md": ROOT / "docs/experiments/审稿意见完成度审计_20260701.md",
    "final_consistency_md": ROOT / "docs/experiments/最终一致性核验_20260701.md",
}
DOCX_FILES = {
    "paper_docx": ROOT / "paper/2026-0268-基于多专家融合的铁路客流多尺度预测方法.docx",
    "response_docx": ROOT / "审稿意见逐条回复稿.docx",
}


def read_docx(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def exists_all(paths: list[str]) -> list[str]:
    return [path for path in paths if not (ROOT / path).exists()]


def has_all_groups(text: str, groups: list[tuple[str, list[str]]]) -> list[str]:
    missing = []
    for label, choices in groups:
        if not any(choice in text for choice in choices):
            missing.append(label)
    return missing


REVIEWER_ITEMS = [
    {
        "id": "R1-1",
        "comment": "统计特征能否解释突发事件、极端天气",
        "files": [
            "docs/experiments/固定消融套件复核_20260701.md",
            "docs/experiments/高需求高波动压力场景审计_20260701.md",
            "docs/experiments/artifacts/event_stress_baseline_audit_20260701.json",
        ],
        "texts": ["paper_md", "response_md", "expert_response_md", "paper_docx", "response_docx"],
        "groups": [
            ("distribution_shift", ["分布偏移", "Distribution Shift"]),
            ("not_causal", ["因果", "不能直接识别", "不声称统计特征能识别"]),
            ("event_stress", ["高需求", "23.9754", "24.8108"]),
        ],
    },
    {
        "id": "R1-2",
        "comment": "数据集介绍增加表格",
        "files": ["paper/论文终稿.md", "paper/2026-0268-基于多专家融合的铁路客流多尺度预测方法.docx"],
        "texts": ["paper_md", "response_md", "expert_response_md", "paper_docx"],
        "groups": [
            ("dataset_table", ["数据集字段", "数据集样例", "表1-0", "表1"]),
            ("real_csv", ["from_nj.csv", "to_nj.csv"]),
        ],
    },
    {
        "id": "R1-3",
        "comment": "补充特征工程及特征选择",
        "files": ["docs/experiments/artifacts/fixed_ablation_metrics_20260701.json"],
        "texts": ["paper_md", "response_md", "expert_response_md", "paper_docx"],
        "groups": [
            ("multi_view", ["multi-view", "多视图", "多视图特征"]),
            ("calendar_features", ["日历场景", "星期", "周末"]),
            ("robust_features", ["robust z-score", "quantile skewness", "peak-to-average", "volatility index"]),
        ],
    },
    {
        "id": "R1-4",
        "comment": "短期专家权重低解释",
        "files": [
            "docs/experiments/Regime门控对齐审计_20260701.md",
            "docs/experiments/场景分层现代强基线审计_20260701.md",
            "docs/experiments/artifacts/regime_gate_alignment_audit_20260701.json",
            "docs/experiments/artifacts/scene_baseline_stratification_audit_20260701.json",
        ],
        "texts": ["paper_md", "response_md", "expert_response_md", "paper_docx", "response_docx"],
        "groups": [
            ("expert_order", ["ShortTerm", "DistributionShift", "专家顺序"]),
            ("weekday_weekend", ["工作日", "周末", "0.5927", "0.5793"]),
            ("regime_alignment", ["0.6714", "0.5909", "峰均比"]),
            ("scene_error", ["18.1153", "19.6382"]),
        ],
    },
    {
        "id": "R1-5",
        "comment": "Traffic 跨粒度泛化跨度大",
        "files": [
            "docs/experiments/Traffic跨粒度证据边界_20260701.md",
            "docs/experiments/Traffic跨粒度边界自动审计_20260701.md",
            "docs/experiments/artifacts/traffic_boundary_audit_20260701.json",
        ],
        "texts": ["paper_md", "response_md", "expert_response_md", "paper_docx", "response_docx"],
        "groups": [
            ("legacy_record", ["现有Traffic补充基准记录", "既有补充基准记录", "既有四步长平均记录"]),
            ("reconstruct", ["重构小时级窗口", "重构时间窗口", "小时级输入窗口"]),
            ("not_main", ["不作为论文主结论证据", "主结论仍以南京南站真实日粒度客流", "主结论和三指标优势只以"]),
            ("mae_boundary", ["MAE不优于", "MAE并非最优", "MAE不是最优"]),
        ],
    },
    {
        "id": "R1-6",
        "comment": "摘要不要出现公式",
        "files": ["paper/论文终稿.md"],
        "texts": ["paper_md", "response_md", "expert_response_md", "paper_docx"],
        "groups": [
            ("textual_abstract", ["摘要", "趋势", "周期", "扰动"]),
            ("formula_removed", ["删除公式", "不出现公式", "文字描述"]),
        ],
    },
    {
        "id": "R1-7",
        "comment": "VAR 表现差，补充/更换对比方法",
        "files": ["docs/experiments/artifacts/modern_baselines_correct_data_20260701.json"],
        "texts": ["paper_md", "response_md", "expert_response_md", "paper_docx"],
        "groups": [
            ("modern_baselines", ["PatchTST", "TimeMixer", "iTransformer", "FreEformer", "DLinear"]),
            ("var_boundary", ["VAR", "不将VAR作为主对比", "不再依赖"]),
        ],
    },
    {
        "id": "R1-8",
        "comment": "英文字体不统一",
        "files": ["docs/experiments/Word格式与表格一致性审计_20260701.md"],
        "texts": ["response_md", "expert_response_md", "completion_md"],
        "groups": [
            ("times", ["Times New Roman"]),
            ("stsong", ["STSong", "宋体"]),
        ],
    },
    {
        "id": "R1-9",
        "comment": "图4第一张图不清晰",
        "files": ["docs/experiments/Word格式与表格一致性审计_20260701.md"],
        "texts": ["paper_md", "response_md", "expert_response_md", "paper_docx"],
        "groups": [
            ("split_figure", ["单独放大", "拆分", "图4(a)", "图4"]),
            ("stl_gate", ["STL", "门控权重"]),
        ],
    },
    {
        "id": "R1-10",
        "comment": "St/Tt/Rt 下标问题",
        "files": ["paper/论文终稿.md"],
        "texts": ["paper_md", "response_md", "expert_response_md", "paper_docx"],
        "groups": [
            ("subscripts", ["Sₜ", "Tₜ", "Rₜ", "S_{t}", "T_{t}", "R_{t}"]),
            ("fixed", ["下标", "统一"]),
        ],
    },
    {
        "id": "R2-1",
        "comment": "创新性不足，像模块组合",
        "files": ["paper/论文终稿.md"],
        "texts": ["paper_md", "response_md", "expert_response_md", "paper_docx"],
        "groups": [
            ("ramr", ["Regime-aware", "RAMR", "场景条件化非平稳分解-路由"]),
            ("not_assembly", ["不在于简单叠加", "不是简单", "模块组合"]),
            ("routing_mapping", ["regime", "expert routing", "专家路由"]),
        ],
    },
    {
        "id": "R2-2",
        "comment": "补充 PatchTST/TimeMixer/iTransformer/FreEformer",
        "files": [
            "docs/experiments/现代强基线配对预测审计_20260701.md",
            "docs/experiments/现代强基线配对显著性审计_20260701.md",
        ],
        "texts": ["paper_md", "response_md", "expert_response_md", "paper_docx"],
        "groups": [
            ("baselines", ["PatchTST", "TimeMixer", "iTransformer", "FreEformer"]),
            ("paired", ["配对预测审计", "Wilcoxon-Holm", "15/15"]),
        ],
    },
    {
        "id": "R2-3",
        "comment": "MSE好但MAE一般，跨域结论不充分",
        "files": [
            "docs/experiments/artifacts/strict_ramr_ve_fixed_ensemble_summary_20260701.json",
            "docs/experiments/现代强基线全指标支配审计_20260701.md",
            "docs/experiments/现代强基线配对显著性审计_20260701.md",
            "docs/experiments/站点方向误差剖面审计_20260701.md",
            "docs/experiments/流量分层误差审计_20260701.md",
        ],
        "texts": ["paper_md", "response_md", "expert_response_md", "paper_docx"],
        "groups": [
            ("final_metrics", ["18.5463", "0.035640", "0.138028"]),
            ("all_metrics", ["15/15", "全指标", "MAE"]),
            ("boundary", ["不能直接外推", "不作为稳定跨域泛化", "边界"]),
        ],
    },
    {
        "id": "R2-4",
        "comment": "增加日历特征、负载均衡消融",
        "files": [
            "docs/experiments/固定消融套件复核_20260701.md",
            "docs/experiments/artifacts/fixed_ablation_metrics_20260701.json",
        ],
        "texts": ["paper_md", "response_md", "expert_response_md", "paper_docx"],
        "groups": [
            ("scene_ablation", ["w/o scene", "无场景门控", "23.3886"]),
            ("regime_ablation", ["w/o regime", "22.9143"]),
            ("balance", ["负载均衡", "方差", "高熵"]),
        ],
    },
    {
        "id": "R2-5",
        "comment": "日粒度表述和熵正则不清晰",
        "files": ["paper/论文终稿.md"],
        "texts": ["paper_md", "response_md", "expert_response_md", "paper_docx"],
        "groups": [
            ("daily_wording", ["相邻日期", "非小时级", "日粒度"]),
            ("entropy_balance", ["高熵负载均衡", "最大化批次平均专家使用熵", "方差型约束仅作为消融"]),
        ],
    },
]


def main() -> int:
    ARTIFACT_JSON.parent.mkdir(parents=True, exist_ok=True)
    texts = {name: path.read_text(encoding="utf-8") for name, path in TEXT_FILES.items()}
    texts.update({name: read_docx(path) for name, path in DOCX_FILES.items()})

    rows: list[dict] = []
    for item in REVIEWER_ITEMS:
        file_missing = exists_all(item["files"])
        combined = "\n".join(texts[name] for name in item["texts"])
        keyword_missing = has_all_groups(combined, item["groups"])
        passed = not file_missing and not keyword_missing
        rows.append(
            {
                "id": item["id"],
                "comment": item["comment"],
                "passed": passed,
                "missing_files": file_missing,
                "missing_keyword_groups": keyword_missing,
                "evidence_files": item["files"],
                "text_scope": item["texts"],
            }
        )

    passed = all(row["passed"] for row in rows)
    report = {
        "passed": passed,
        "protocol": (
            "Read-only reviewer completion audit. Checks evidence files and text coverage for all "
            "15 reviewer comments; performs no training, tuning, checkpoint selection, or test-set evaluation."
        ),
        "n_items": len(rows),
        "passed_items": sum(1 for row in rows if row["passed"]),
        "rows": rows,
    }
    ARTIFACT_JSON.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        "# 审稿意见完成度自动审计（2026-07-01）",
        "",
        f"总体结果：{'通过' if passed else '未通过'}",
        "",
        "## 审计口径",
        "",
        "- 本审计覆盖专家意见1的10条意见和专家意见2的5条意见，共15条。",
        "- 审计只检查证据文件和文稿表述，不训练模型、不调参、不重新评估测试集。",
        "- 对每条意见同时检查至少一个可追溯证据文件和论文/回复中的关键表述。",
        "",
        "## 覆盖矩阵",
        "",
        "| 编号 | 意见 | 结果 | 缺失文件 | 缺失关键词组 |",
        "| --- | --- | --- | --- | --- |",
    ]
    for row in rows:
        missing_files = "<br>".join(row["missing_files"]) if row["missing_files"] else "-"
        missing_keywords = ", ".join(row["missing_keyword_groups"]) if row["missing_keyword_groups"] else "-"
        lines.append(
            f"| {row['id']} | {row['comment']} | {'通过' if row['passed'] else '失败'} | "
            f"{missing_files} | {missing_keywords} |"
        )

    lines.extend(
        [
            "",
            "## 结论",
            "",
            f"- 已通过条目：{sum(1 for row in rows if row['passed'])}/{len(rows)}。",
            "- 审计通过表示每条意见均有证据文件和论文/回复文字闭环；不表示所有外部数据集均已端到端复现。",
            "- Traffic 仍按边界审计处理：当前仓库无可端到端重跑的 Traffic 原始数据，相关结果只作为既有补充基准记录。",
        ]
    )
    REPORT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if passed else 1


if __name__ == "__main__":
    raise SystemExit(main())
