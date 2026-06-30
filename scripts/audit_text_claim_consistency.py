#!/usr/bin/env python3
"""Audit numeric and wording consistency across the revised paper package.

The script is intentionally read-only with respect to model outputs: it only
loads fixed artifacts and manuscript/reply text, then verifies that the claims
written in the paper match the evidence files. No training, checkpoint
selection, calibration, or test-set re-evaluation is performed here.
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
ARTIFACT_DIR = ROOT / "docs" / "experiments" / "artifacts"
REPORT_MD = ROOT / "docs" / "experiments" / "论文数值与口径一致性审计_20260701.md"
REPORT_JSON = ARTIFACT_DIR / "text_claim_consistency_audit_20260701.json"


TEXT_FILES = {
    "readme": ROOT / "README.md",
    "paper_md": ROOT / "paper" / "论文终稿.md",
    "response_md": ROOT / "审稿意见逐条回复稿.md",
    "expert_response_md": ROOT / "专家意见逐条修改说明.md",
    "final_consistency_md": ROOT / "docs" / "experiments" / "最终一致性核验_20260701.md",
    "paired_audit_md": ROOT / "docs" / "experiments" / "现代强基线配对预测审计_20260701.md",
}

DOCX_FILES = {
    "paper_docx": ROOT / "paper" / "2026-0268-基于多专家融合的铁路客流多尺度预测方法.docx",
    "response_docx": ROOT / "审稿意见逐条回复稿.docx",
}

EXPECTED_HASHES = {
    "from_nj": "c4f8919ce83c4c677d772036ed42fb5e76c99507f3368e2d32f081b556df4ae7",
    "to_nj": "d16b3becfc12ed22a576741a505080e9dc279d4ab54d8c8ef597b729efb3c41e",
}

BANNED_OVERCLAIMS = [
    "跨域泛化能力较强",
    "稳定跨域泛化能力",
    "直接跨分辨率迁移",
    "逐样本配对显著性检验",
]

NEGATION_MARKERS = [
    "不",
    "不能",
    "不得",
    "不宜",
    "不主张",
    "不能声称",
    "尚不足以证明",
    "并非",
    "只作为",
]


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def read_docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def add_check(checks: list[dict], name: str, passed: bool, detail: dict | list | str) -> None:
    checks.append({"name": name, "passed": bool(passed), "detail": detail})


def contains_all(text: str, needles: list[str]) -> list[str]:
    return [needle for needle in needles if needle not in text]


def missing_groups(text: str, groups: list[tuple[str, list[str]]]) -> list[str]:
    missing = []
    for label, alternatives in groups:
        if not any(item in text for item in alternatives):
            missing.append(label)
    return missing


def non_negated_hits(text: str, phrases: list[str]) -> list[str]:
    hits = []
    for line in text.splitlines():
        for phrase in phrases:
            if phrase not in line:
                continue
            if any(marker in line for marker in NEGATION_MARKERS):
                continue
            hits.append(phrase)
    return sorted(set(hits))


def main() -> int:
    REPORT_JSON.parent.mkdir(parents=True, exist_ok=True)
    checks: list[dict] = []

    texts = {name: path.read_text(encoding="utf-8") for name, path in TEXT_FILES.items()}
    texts.update({name: read_docx_text(path) for name, path in DOCX_FILES.items()})

    final_summary = read_json(ARTIFACT_DIR / "strict_ramr_ve_fixed_ensemble_summary_20260701.json")
    no_leak = read_json(ARTIFACT_DIR / "no_leakage_protocol_audit_20260701.json")
    dominance = read_json(ARTIFACT_DIR / "modern_baseline_dominance_audit_20260701.json")
    paired = read_json(ARTIFACT_DIR / "modern_baseline_paired_audit_20260701.json")
    paired_sig = read_json(ARTIFACT_DIR / "paired_significance_audit_20260701.json")
    event_stress = read_json(ARTIFACT_DIR / "event_stress_baseline_audit_20260701.json")
    scene_baseline = read_json(ARTIFACT_DIR / "scene_baseline_stratification_audit_20260701.json")
    target_profile = read_json(ARTIFACT_DIR / "target_error_profile_20260701.json")
    flow = read_json(ARTIFACT_DIR / "flow_stratified_error_audit_20260701.json")
    regime = read_json(ARTIFACT_DIR / "regime_gate_alignment_audit_20260701.json")

    metric_detail = {
        "mape": final_summary["test"]["metrics"]["mape"],
        "mse": final_summary["test"]["metrics"]["mse"],
        "mae": final_summary["test"]["metrics"]["mae"],
        "valid_count": final_summary["test"]["metrics"]["valid_count"],
    }
    add_check(
        checks,
        "authoritative_final_metrics_loaded",
        round(metric_detail["mape"], 4) == 18.5463
        and round(metric_detail["mse"], 6) == 0.035640
        and round(metric_detail["mae"], 6) == 0.138028
        and metric_detail["valid_count"] == 4240,
        metric_detail,
    )

    no_leak_ok = (
        no_leak["passed"] is True
        and no_leak["n_days"] == 846
        and no_leak["n_targets"] == 20
        and no_leak["split"]["train"]["n_days"] == 423
        and no_leak["split"]["validation"]["n_days"] == 211
        and no_leak["split"]["test"]["n_days"] == 212
        and no_leak["checks"]["robust_scaler_matches_train_only_fit"] is True
        and no_leak["checks"]["robust_scaler_does_not_match_full_data_fit"] is True
        and no_leak["data_hash"]["from_nj_sha256"] == EXPECTED_HASHES["from_nj"]
        and no_leak["data_hash"]["to_nj_sha256"] == EXPECTED_HASHES["to_nj"]
    )
    add_check(
        checks,
        "authoritative_no_leakage_artifact_loaded",
        no_leak_ok,
        {
            "split": no_leak["split"],
            "from_hash": no_leak["data_hash"]["from_nj_sha256"],
            "to_hash": no_leak["data_hash"]["to_nj_sha256"],
            "scaler_checks": {
                "train_only": no_leak["checks"]["robust_scaler_matches_train_only_fit"],
                "not_full_data": no_leak["checks"]["robust_scaler_does_not_match_full_data_fit"],
            },
        },
    )

    dominance_ok = (
        dominance["point_dominance_cells"] == 15
        and dominance["point_dominance_total_cells"] == 15
        and dominance["all_baselines_all_three_point_better"] is True
    )
    add_check(
        checks,
        "authoritative_dominance_artifact_loaded",
        dominance_ok,
        {
            "point_dominance_cells": dominance["point_dominance_cells"],
            "point_dominance_total_cells": dominance["point_dominance_total_cells"],
            "closest_baseline": dominance["closest_baseline_by_minimum_relative_reduction"],
        },
    )

    paired_ok = (
        paired["passed"] is True
        and paired["n_rows"] == 212
        and paired["n_targets"] == 20
        and paired["summary"]["point_dominance_cells"] == 15
        and paired["summary"]["point_dominance_total_cells"] == 15
        and round(paired["summary"]["min_all_three_prob"], 3) == 0.894
        and paired["summary"]["closest_by_all_three_prob"] == "TimeMixer"
    )
    add_check(
        checks,
        "authoritative_paired_artifact_loaded",
        paired_ok,
        {
            "point_dominance": f"{paired['summary']['point_dominance_cells']}/{paired['summary']['point_dominance_total_cells']}",
            "min_all_three_prob": paired["summary"]["min_all_three_prob"],
            "closest": paired["summary"]["closest_by_all_three_prob"],
        },
    )

    paired_sig_summary = paired_sig["summary"]
    paired_sig_ok = (
        paired_sig["passed"] is True
        and paired_sig["n_dates"] == 212
        and paired_sig["n_targets"] == 20
        and paired_sig_summary["point_margin_positive_cells"] == 15
        and paired_sig_summary["total_cells"] == 15
        and paired_sig_summary["wilcoxon_holm_significant_0_01_cells"] == 15
        and paired_sig_summary["bootstrap_ci95_low_positive_cells"] == 14
        and paired_sig_summary["sign_flip_holm_significant_0_01_cells"] == 14
        and paired_sig_summary["closest_bootstrap_ci_cell"]["baseline"] == "TimeMixer"
        and paired_sig_summary["closest_bootstrap_ci_cell"]["metric"] == "MSE"
    )
    add_check(
        checks,
        "authoritative_paired_significance_artifact_loaded",
        paired_sig_ok,
        {
            "point_margin_positive_cells": paired_sig_summary["point_margin_positive_cells"],
            "wilcoxon_holm_significant_0_01_cells": paired_sig_summary["wilcoxon_holm_significant_0_01_cells"],
            "bootstrap_ci95_low_positive_cells": paired_sig_summary["bootstrap_ci95_low_positive_cells"],
            "sign_flip_holm_significant_0_01_cells": paired_sig_summary["sign_flip_holm_significant_0_01_cells"],
            "closest_boundary": paired_sig_summary["closest_bootstrap_ci_cell"],
        },
    )

    event_summary = event_stress["summary"]
    event_stress_ok = (
        event_stress["passed"] is True
        and event_stress["n_dates"] == 212
        and event_stress["n_targets"] == 20
        and event_stress["test_start_index"] == 634
        and event_summary["high_demand_ramr_best_all_metrics"] is True
        and event_summary["high_abs_shift_ramr_best_mape"] is True
        and event_summary["high_abs_shift_best_mse_model"] == "TimeMixer"
        and event_summary["high_abs_shift_best_mae_model"] == "TimeMixer"
    )
    add_check(
        checks,
        "authoritative_event_stress_artifact_loaded",
        event_stress_ok,
        {
            "high_demand_ramr_best_all_metrics": event_summary["high_demand_ramr_best_all_metrics"],
            "high_abs_shift_ramr_best_mape": event_summary["high_abs_shift_ramr_best_mape"],
            "high_abs_shift_best_mse_model": event_summary["high_abs_shift_best_mse_model"],
            "high_abs_shift_best_mae_model": event_summary["high_abs_shift_best_mae_model"],
        },
    )

    scene_summary = scene_baseline["summary"]
    scene_ok = (
        scene_baseline["passed"] is True
        and scene_baseline["n_dates"] == 212
        and scene_baseline["n_targets"] == 20
        and scene_baseline["test_start_index"] == 634
        and scene_baseline["scene_counts"]["weekday"] == 152
        and scene_baseline["scene_counts"]["weekend"] == 60
        and scene_baseline["scene_counts"]["holiday"] == 0
        and scene_summary["weekday_ramr_best_all_metrics"] is True
        and scene_summary["weekend_ramr_best_all_metrics"] is True
        and round(scene_summary["weekday_ramr_metrics"]["mape"], 4) == 18.1153
        and round(scene_summary["weekend_ramr_metrics"]["mape"], 4) == 19.6382
    )
    add_check(
        checks,
        "authoritative_scene_baseline_artifact_loaded",
        scene_ok,
        {
            "weekday_days": scene_baseline["scene_counts"]["weekday"],
            "weekend_days": scene_baseline["scene_counts"]["weekend"],
            "holiday_days": scene_baseline["scene_counts"]["holiday"],
            "weekday_mape": scene_summary["weekday_ramr_metrics"]["mape"],
            "weekend_mape": scene_summary["weekend_ramr_metrics"]["mape"],
        },
    )

    core_metric_needles = ["18.5463", "0.035640", "0.138028"]
    core_docs = ["readme", "paper_md", "response_md", "expert_response_md", "final_consistency_md", "paper_docx"]
    metric_missing = {
        name: contains_all(texts[name], core_metric_needles) for name in core_docs
    }
    metric_missing = {name: miss for name, miss in metric_missing.items() if miss}
    add_check(checks, "final_metrics_written_in_core_texts", not metric_missing, metric_missing)

    no_leak_groups = [
        ("423/211/212", ["423/211/212", "训练集 423、验证集 211、测试集 212"]),
        ("RobustScaler", ["RobustScaler"]),
        ("训练集拟合", ["训练集拟合", "训练集拟合 Scaler", "仅在训练集拟合"]),
    ]
    no_leak_docs = ["paper_md", "response_md", "expert_response_md", "final_consistency_md", "paper_docx", "response_docx"]
    no_leak_missing = {
        name: missing_groups(texts[name], no_leak_groups) for name in no_leak_docs
    }
    no_leak_missing = {name: miss for name, miss in no_leak_missing.items() if miss}
    add_check(checks, "no_leakage_protocol_written_in_texts", not no_leak_missing, no_leak_missing)

    paired_groups = [
        ("配对预测审计", ["配对预测审计"]),
        ("15/15", ["15/15", "15个点估计比较单元", "15 个点估计比较单元"]),
        ("89.40", ["89.40"]),
        ("TimeMixer", ["TimeMixer"]),
    ]
    paired_docs = ["readme", "paper_md", "response_md", "expert_response_md", "final_consistency_md", "paper_docx", "response_docx"]
    paired_missing = {
        name: missing_groups(texts[name], paired_groups) for name in paired_docs
    }
    paired_missing = {name: miss for name, miss in paired_missing.items() if miss}
    add_check(checks, "paired_modern_baseline_claims_written_in_texts", not paired_missing, paired_missing)

    paired_sig_groups = [
        ("Wilcoxon-Holm", ["Wilcoxon-Holm"]),
        ("15/15", ["15/15"]),
        ("14/15", ["14/15"]),
        ("TimeMixer", ["TimeMixer"]),
        ("MSE", ["MSE"]),
    ]
    paired_sig_docs = ["readme", "paper_md", "response_md", "expert_response_md", "final_consistency_md", "paper_docx", "response_docx"]
    paired_sig_missing = {
        name: missing_groups(texts[name], paired_sig_groups) for name in paired_sig_docs
    }
    paired_sig_missing = {name: miss for name, miss in paired_sig_missing.items() if miss}
    add_check(checks, "paired_significance_claims_written_in_texts", not paired_sig_missing, paired_sig_missing)

    event_groups = [
        ("高需求", ["高需求"]),
        ("23.9754", ["23.9754"]),
        ("高日际变化", ["高日际变化", "日际客流绝对变化", "日际突变"]),
        ("24.8108", ["24.8108"]),
        ("TimeMixer", ["TimeMixer"]),
        ("因果", ["因果"]),
    ]
    event_docs = ["readme", "paper_md", "response_md", "expert_response_md", "final_consistency_md", "paper_docx", "response_docx"]
    event_missing = {
        name: missing_groups(texts[name], event_groups) for name in event_docs
    }
    event_missing = {name: miss for name, miss in event_missing.items() if miss}
    add_check(checks, "event_stress_claims_written_in_texts", not event_missing, event_missing)

    scene_groups = [
        ("场景分层", ["场景分层"]),
        ("工作日", ["工作日"]),
        ("周末", ["周末"]),
        ("18.1153", ["18.1153"]),
        ("19.6382", ["19.6382"]),
        ("节假日", ["节假日"]),
    ]
    scene_docs = ["readme", "paper_md", "response_md", "expert_response_md", "final_consistency_md", "paper_docx", "response_docx"]
    scene_missing = {
        name: missing_groups(texts[name], scene_groups) for name in scene_docs
    }
    scene_missing = {name: miss for name, miss in scene_missing.items() if miss}
    add_check(checks, "scene_baseline_claims_written_in_texts", not scene_missing, scene_missing)

    boundary_groups = [
        ("结构鲁棒", ["结构鲁棒"]),
        ("不能证明", ["不能证明", "不证明", "尚不足以证明", "不足以证明", "不将单一", "不能直接外推"]),
        ("不替换", ["不替换"]),
    ]
    boundary_docs = ["paper_md", "response_md", "expert_response_md", "final_consistency_md", "paper_docx", "response_docx"]
    boundary_missing = {
        name: missing_groups(texts[name], boundary_groups) for name in boundary_docs
    }
    boundary_missing = {name: miss for name, miss in boundary_missing.items() if miss}
    add_check(checks, "boundary_language_written_in_texts", not boundary_missing, boundary_missing)

    stale_hits = {}
    for name in ["readme", "paper_md", "expert_response_md", "final_consistency_md", "paper_docx"]:
        hits = non_negated_hits(texts[name], BANNED_OVERCLAIMS)
        if hits:
            stale_hits[name] = hits
    add_check(checks, "overclaim_phrase_scan", not stale_hits, stale_hits)

    profile_ok = (
        target_profile["n_rows"] == 212
        and target_profile["n_targets"] == 20
        and target_profile["summary"]["targets_below_20_mape"] == 16
        and target_profile["summary"]["targets_below_30_mape"] == 17
    )
    add_check(
        checks,
        "target_profile_artifact_loaded",
        profile_ok,
        {
            "targets_below_20_mape": target_profile["summary"]["targets_below_20_mape"],
            "targets_below_30_mape": target_profile["summary"]["targets_below_30_mape"],
        },
    )

    flow_groups = {(row["group_type"], row["group"]): row for row in flow["rows"]}
    flow_ok = (
        flow["n_rows"] == 212
        and flow_groups[("flow_magnitude", "high_flow_>=3000")]["mape"] < 15.0
        and flow_groups[("daily_total", "high_demand_days")]["mape"]
        > flow_groups[("daily_total", "normal_demand_days")]["mape"]
    )
    add_check(
        checks,
        "flow_stratified_artifact_loaded",
        flow_ok,
        {
            "high_flow_mape": flow_groups[("flow_magnitude", "high_flow_>=3000")]["mape"],
            "normal_demand_mape": flow_groups[("daily_total", "normal_demand_days")]["mape"],
            "high_demand_mape": flow_groups[("daily_total", "high_demand_days")]["mape"],
        },
    )

    regime_ok = (
        regime["expert_order"] == ["ShortTerm", "LongTerm", "DistributionShift"]
        and regime["interpretation"]["high_daily_vol_short_delta"] > 0.20
        and regime["interpretation"]["high_peak_distribution_delta"] > 0.20
    )
    add_check(
        checks,
        "regime_gate_artifact_loaded",
        regime_ok,
        {
            "high_daily_vol_short_delta": regime["interpretation"]["high_daily_vol_short_delta"],
            "high_peak_distribution_delta": regime["interpretation"]["high_peak_distribution_delta"],
        },
    )

    passed = all(item["passed"] for item in checks)
    report = {
        "passed": passed,
        "protocol": "Text-claim audit only. Reads fixed evidence artifacts and manuscript/reply text; performs no training, tuning, checkpoint selection, or test-set re-evaluation.",
        "checks": checks,
    }
    REPORT_JSON.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        "# 论文数值与口径一致性审计（2026-07-01）",
        "",
        f"总体结果：{'通过' if passed else '未通过'}",
        "",
        "## 审计原则",
        "",
        "本审计只读取固定证据文件和文稿文本，不训练模型、不调参、不重新选择 checkpoint、不重新评估候选测试集结果。其目的不是生成新指标，而是确认论文、审稿回复和逐条说明中的数字与实验边界均来自同一组真实 artifact。",
        "",
        "## 权威证据",
        "",
        f"- 最终 Strict RAMR-VE：MAPE={metric_detail['mape']:.4f}%、MSE={metric_detail['mse']:.6f}、MAE={metric_detail['mae']:.6f}。",
        f"- 数据与协议：846 个共同日期、20 个目标变量、423/211/212 天时间顺序划分，RobustScaler 仅在训练集拟合。",
        f"- 现代强基线点估计支配：{dominance['point_dominance_cells']}/{dominance['point_dominance_total_cells']} 个 MAPE/MSE/MAE 比较单元更低。",
        f"- 现代强基线配对预测审计：{paired['summary']['point_dominance_cells']}/{paired['summary']['point_dominance_total_cells']} 个点估计比较单元更低，三指标同时更优的日期级 bootstrap 概率最低为 {paired['summary']['min_all_three_prob'] * 100:.2f}%，最接近边界为 {paired['summary']['closest_by_all_three_prob']}。",
        f"- 现代强基线配对统计审计：{paired_sig_summary['wilcoxon_holm_significant_0_01_cells']}/{paired_sig_summary['total_cells']} 个 Wilcoxon-Holm 单侧检验在 α=0.01 下显著，bootstrap/符号置换边界单元为 TimeMixer-MSE。",
        f"- 高需求/高波动压力审计：高需求日 Strict RAMR-VE 三指标均优于现代强基线；高日际变化日 MAPE 最优，但 MSE/MAE 边界模型为 TimeMixer。",
        f"- 场景分层现代强基线审计：工作日 {scene_baseline['scene_counts']['weekday']} 天、周末 {scene_baseline['scene_counts']['weekend']} 天上 Strict RAMR-VE 三指标均最优；节假日样本数为 {scene_baseline['scene_counts']['holiday']}。",
        "",
        "## 检查结果",
        "",
        "| 检查项 | 结果 | 摘要 |",
        "| --- | --- | --- |",
    ]
    for item in checks:
        detail = item["detail"]
        if isinstance(detail, dict):
            summary = ", ".join(f"{k}={v}" for k, v in list(detail.items())[:4])
        elif isinstance(detail, list):
            summary = f"{len(detail)} item(s)"
        else:
            summary = str(detail)
        lines.append(f"| {item['name']} | {'通过' if item['passed'] else '失败'} | {summary} |")

    lines.extend(
        [
            "",
            "## 结论",
            "",
            "- 论文正文、审稿回复、逐条修改说明、最终一致性核验和 Word 稿均已写入最终三指标、无泄露协议、现代强基线全指标支配、配对预测审计、配对统计审计、压力场景审计和场景分层审计证据。",
            "- 文稿保留了必要边界：Traffic 只作为重构小时级窗口和周期特征后的结构鲁棒性初步验证，不作为稳定跨域泛化证明；配对预测审计作为补充证据，不替换原现代强基线主表。",
            "- 当前结论限定在南京南站双向日粒度客流、当前时间顺序划分和当前对比模型集合下。",
        ]
    )
    REPORT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")

    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if passed else 1


if __name__ == "__main__":
    sys.exit(main())
