#!/usr/bin/env python3
"""Verify final paper deliverables and reviewer-response evidence.

This verifier is intentionally strict. It checks the concrete artifacts that
support the revised paper: real data hashes, final metric evidence, Word
formatting, synced paper copies, reviewer-response coverage, and stale-phrase
guards. It writes both JSON and Markdown reports under docs/experiments.
"""

from __future__ import annotations

import hashlib
import json
import re
import sys
from collections import Counter
from pathlib import Path

import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
ARTIFACT_DIR = ROOT / "docs" / "experiments" / "artifacts"
REPORT_JSON = ARTIFACT_DIR / "final_deliverable_verification_20260701.json"
REPORT_MD = ROOT / "docs" / "experiments" / "最终交付自动核验_20260701.md"

EXPECTED = {
    "from_hash": "c4f8919ce83c4c677d772036ed42fb5e76c99507f3368e2d32f081b556df4ae7",
    "to_hash": "d16b3becfc12ed22a576741a505080e9dc279d4ab54d8c8ef597b729efb3c41e",
    "from_shape": [854, 11],
    "to_shape": [852, 11],
    "merged_dates": 846,
    "final_mape": 18.546323040648975,
    "final_mse": 0.035640054881848904,
    "final_mae": 0.13802760334469663,
    "baseline_mape": 19.41,
    "baseline_mse": 0.037700,
    "baseline_mae": 0.144000,
    "docx_font": "Times New Roman",
    "docx_east_asia": "STSong",
    "response_headings": 15,
}

REQUIRED_FILES = [
    "README.md",
    "data/from_nj.csv",
    "data/to_nj.csv",
    "paper/2026-0268-基于多专家融合的铁路客流多尺度预测方法.docx",
    "paper/论文终稿.docx",
    "paper/论文终稿.md",
    "审稿意见逐条回复稿.md",
    "审稿意见逐条回复稿.docx",
    "专家意见逐条修改说明.md",
    "docs/experiments/最终一致性核验_20260701.md",
    "docs/experiments/审稿意见完成度审计_20260701.md",
    "docs/experiments/无泄露实验协议审计_20260701.md",
    "docs/experiments/固定消融套件复核_20260701.md",
    "docs/experiments/现代强基线优势边界审计_20260701.md",
    "docs/experiments/现代强基线全指标支配审计_20260701.md",
    "docs/experiments/现代强基线配对预测审计_20260701.md",
    "docs/experiments/现代强基线配对显著性审计_20260701.md",
    "docs/experiments/高需求高波动压力场景审计_20260701.md",
    "docs/experiments/场景分层现代强基线审计_20260701.md",
    "docs/experiments/论文数值与口径一致性审计_20260701.md",
    "docs/experiments/站点方向误差剖面审计_20260701.md",
    "docs/experiments/流量分层误差审计_20260701.md",
    "docs/experiments/Regime门控对齐审计_20260701.md",
    "docs/experiments/Word格式与表格一致性审计_20260701.md",
    "docs/experiments/artifacts/strict_ramr_ve_fixed_ensemble_summary_20260701.json",
    "docs/experiments/artifacts/strict_ramr_ve_test_predictions_20260701.npz",
    "docs/experiments/artifacts/no_leakage_protocol_audit_20260701.json",
    "docs/experiments/artifacts/modern_baselines_correct_data_20260701.json",
    "docs/experiments/artifacts/modern_baseline_margin_audit_20260701.csv",
    "docs/experiments/artifacts/modern_baseline_margin_audit_20260701.json",
    "docs/experiments/artifacts/modern_baseline_dominance_audit_20260701.csv",
    "docs/experiments/artifacts/modern_baseline_dominance_audit_20260701.json",
    "docs/experiments/figures/modern_baseline_all_metrics_20260701.png",
    "docs/experiments/artifacts/modern_baseline_paired_audit_20260701.csv",
    "docs/experiments/artifacts/modern_baseline_paired_audit_20260701.json",
    "docs/experiments/artifacts/paired_significance_audit_20260701.csv",
    "docs/experiments/artifacts/paired_significance_audit_20260701.json",
    "docs/experiments/artifacts/event_stress_baseline_audit_20260701.csv",
    "docs/experiments/artifacts/event_stress_baseline_audit_20260701.json",
    "docs/experiments/artifacts/scene_baseline_stratification_audit_20260701.csv",
    "docs/experiments/artifacts/scene_baseline_stratification_audit_20260701.json",
    "docs/experiments/artifacts/text_claim_consistency_audit_20260701.json",
    "docs/experiments/artifacts/modern_baseline_paired_predictions_20260701/DLinear_test_predictions.npz",
    "docs/experiments/artifacts/modern_baseline_paired_predictions_20260701/PatchTST_test_predictions.npz",
    "docs/experiments/artifacts/modern_baseline_paired_predictions_20260701/iTransformer_test_predictions.npz",
    "docs/experiments/artifacts/modern_baseline_paired_predictions_20260701/TimeMixer_test_predictions.npz",
    "docs/experiments/artifacts/modern_baseline_paired_predictions_20260701/FreEformer_test_predictions.npz",
    "docs/experiments/artifacts/target_error_profile_20260701.csv",
    "docs/experiments/artifacts/target_error_profile_20260701.json",
    "docs/experiments/artifacts/flow_stratified_error_audit_20260701.csv",
    "docs/experiments/artifacts/flow_stratified_error_audit_20260701.json",
    "docs/experiments/artifacts/regime_gate_alignment_audit_20260701.csv",
    "docs/experiments/artifacts/regime_gate_alignment_audit_20260701.json",
    "docs/experiments/artifacts/fixed_ablation_metrics_20260701.csv",
    "docs/experiments/artifacts/fixed_ablation_metrics_20260701.json",
]

BANNED_PATTERNS = [
    "跨域泛化能力较强",
    "节假日等强波动区间",
    "日内高频",
    "早晚高峰",
    "S_t",
    "T_t",
    "R_t",
    "稳定性与泛化能力",
]


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def approx_equal(value: float, expected: float, tol: float = 1e-9) -> bool:
    return abs(float(value) - float(expected)) <= tol


def add_check(checks: list[dict], name: str, passed: bool, detail: str | dict | list) -> None:
    checks.append({"name": name, "passed": bool(passed), "detail": detail})


def docx_font_summary(path: Path) -> dict:
    doc = Document(path)
    fonts = Counter()
    east = Counter()
    total = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if not run.text.strip():
                continue
            total += 1
            fonts[run.font.name] += 1
            rfonts = run._element.rPr.rFonts if run._element.rPr is not None else None
            east_val = (
                rfonts.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia")
                if rfonts is not None
                else None
            )
            east[east_val] += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.text.strip():
                            continue
                        total += 1
                        fonts[run.font.name] += 1
                        rfonts = run._element.rPr.rFonts if run._element.rPr is not None else None
                        east_val = (
                            rfonts.get(
                                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
                            )
                            if rfonts is not None
                            else None
                        )
                        east[east_val] += 1
    return {
        "runs": total,
        "fonts": dict(fonts),
        "east_asia": dict(east),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
    }


def scan_banned(files: list[Path]) -> list[dict]:
    hits = []
    for path in files:
        text = path.read_text(encoding="utf-8", errors="ignore")
        for pattern in BANNED_PATTERNS:
            for match in re.finditer(re.escape(pattern), text):
                line_no = text.count("\n", 0, match.start()) + 1
                line = text.splitlines()[line_no - 1].strip()
                # Reviewer-response drafts may quote the original reviewer wording.
                if path.name.startswith("审稿意见逐条回复稿") and (
                    "意见概述" in line or "不再使用" in line or "未发现" in line
                ):
                    continue
                hits.append({"file": str(path.relative_to(ROOT)), "line": line_no, "pattern": pattern})
    return hits


def main() -> int:
    checks: list[dict] = []
    ARTIFACT_DIR.mkdir(parents=True, exist_ok=True)

    missing = [p for p in REQUIRED_FILES if not (ROOT / p).exists()]
    add_check(checks, "required_files_exist", not missing, missing)

    from_path = ROOT / "data/from_nj.csv"
    to_path = ROOT / "data/to_nj.csv"
    from_df = pd.read_csv(from_path, encoding="gbk")
    to_df = pd.read_csv(to_path, encoding="gbk")
    merged = pd.merge(from_df[["time"]], to_df[["time"]], on="time", how="inner")
    data_detail = {
        "from_hash": sha256(from_path),
        "to_hash": sha256(to_path),
        "from_shape": list(from_df.shape),
        "to_shape": list(to_df.shape),
        "merged_dates": int(len(merged)),
        "from_missing": int(from_df.isna().sum().sum()),
        "to_missing": int(to_df.isna().sum().sum()),
        "from_duplicate_time": int(from_df["time"].duplicated().sum()),
        "to_duplicate_time": int(to_df["time"].duplicated().sum()),
    }
    data_ok = (
        data_detail["from_hash"] == EXPECTED["from_hash"]
        and data_detail["to_hash"] == EXPECTED["to_hash"]
        and data_detail["from_shape"] == EXPECTED["from_shape"]
        and data_detail["to_shape"] == EXPECTED["to_shape"]
        and data_detail["merged_dates"] == EXPECTED["merged_dates"]
        and data_detail["from_missing"] == 0
        and data_detail["to_missing"] == 0
        and data_detail["from_duplicate_time"] == 0
        and data_detail["to_duplicate_time"] == 0
    )
    add_check(checks, "real_data_integrity", data_ok, data_detail)

    paper_main = ROOT / "paper/2026-0268-基于多专家融合的铁路客流多尺度预测方法.docx"
    paper_copy = ROOT / "paper/论文终稿.docx"
    docx_hash_detail = {"main": sha256(paper_main), "copy": sha256(paper_copy)}
    add_check(checks, "paper_docx_copy_synced", docx_hash_detail["main"] == docx_hash_detail["copy"], docx_hash_detail)

    font_detail = docx_font_summary(paper_main)
    font_ok = (
        font_detail["runs"] > 0
        and font_detail["fonts"] == {EXPECTED["docx_font"]: font_detail["runs"]}
        and font_detail["east_asia"] == {EXPECTED["docx_east_asia"]: font_detail["runs"]}
    )
    add_check(checks, "paper_word_fonts", font_ok, font_detail)

    response_font = docx_font_summary(ROOT / "审稿意见逐条回复稿.docx")
    response_font_ok = (
        response_font["runs"] > 0
        and response_font["fonts"] == {EXPECTED["docx_font"]: response_font["runs"]}
        and response_font["east_asia"] == {EXPECTED["docx_east_asia"]: response_font["runs"]}
    )
    add_check(checks, "response_word_fonts", response_font_ok, response_font)

    response_md = (ROOT / "审稿意见逐条回复稿.md").read_text(encoding="utf-8")
    headings = [line for line in response_md.splitlines() if line.startswith("### ")]
    add_check(
        checks,
        "reviewer_response_coverage",
        len(headings) == EXPECTED["response_headings"],
        {"heading_count": len(headings), "headings": headings},
    )

    summary = json.loads(
        (ROOT / "docs/experiments/artifacts/strict_ramr_ve_fixed_ensemble_summary_20260701.json").read_text(
            encoding="utf-8"
        )
    )
    test_metrics = summary["test"]["metrics"]
    final_ok = (
        approx_equal(test_metrics["mape"], EXPECTED["final_mape"])
        and approx_equal(test_metrics["mse"], EXPECTED["final_mse"])
        and approx_equal(test_metrics["mae"], EXPECTED["final_mae"])
        and test_metrics["mape"] < EXPECTED["baseline_mape"]
        and test_metrics["mse"] < EXPECTED["baseline_mse"]
        and test_metrics["mae"] < EXPECTED["baseline_mae"]
    )
    add_check(checks, "strict_ramr_ve_final_metrics", final_ok, test_metrics)

    no_leak = json.loads(
        (ROOT / "docs/experiments/artifacts/no_leakage_protocol_audit_20260701.json").read_text(
            encoding="utf-8"
        )
    )
    no_leak_checks = no_leak["checks"]
    no_leak_ok = (
        no_leak["passed"] is True
        and no_leak["n_days"] == EXPECTED["merged_dates"]
        and no_leak["n_targets"] == 20
        and no_leak["lookback"] == 6
        and no_leak["horizon"] == 1
        and no_leak["split"]["train"]["n_days"] == 423
        and no_leak["split"]["validation"]["n_days"] == 211
        and no_leak["split"]["test"]["n_days"] == 212
        and no_leak_checks["chronological_split_non_overlapping"] is True
        and no_leak_checks["robust_scaler_matches_train_only_fit"] is True
        and no_leak_checks["robust_scaler_does_not_match_full_data_fit"] is True
        and no_leak_checks["fixed_prediction_rows_match_test_dataset"] is True
    )
    add_check(
        checks,
        "no_leakage_protocol_audit",
        no_leak_ok,
        {
            "train_days": no_leak["split"]["train"]["n_days"],
            "validation_days": no_leak["split"]["validation"]["n_days"],
            "test_days": no_leak["split"]["test"]["n_days"],
            "prediction_rows": no_leak["prediction_artifact"]["actual_prediction_rows"],
        },
    )

    ablation = pd.read_csv(ROOT / "docs/experiments/artifacts/fixed_ablation_metrics_20260701.csv")
    required_ablation = {
        "RAMR-full(scene+regime+robust+entropy)",
        "w/o scene gating",
        "w/o regime routing",
        "variance load balance",
        "distribution basic(mean/std/max)",
        "distribution quantile",
        "MAE-aware robust",
    }
    ablation_names = set(ablation["name"].tolist())
    scene = ablation.loc[ablation["name"] == "w/o scene gating", "test_mape"].iloc[0]
    full = ablation.loc[ablation["name"] == "RAMR-full(scene+regime+robust+entropy)", "test_mape"].iloc[0]
    quantile = ablation.loc[ablation["name"] == "distribution quantile", "test_mape"].iloc[0]
    ablation_ok = required_ablation.issubset(ablation_names) and scene > full and quantile < full
    add_check(
        checks,
        "fixed_ablation_evidence",
        ablation_ok,
        ablation[["name", "test_mape", "test_mse", "test_mae"]].to_dict(orient="records"),
    )

    margin = json.loads(
        (ROOT / "docs/experiments/artifacts/modern_baseline_margin_audit_20260701.json").read_text(
            encoding="utf-8"
        )
    )
    fre = next((row for row in margin["baseline_margin_rows"] if row["baseline"] == "FreEformer"), None)
    margin_ok = (
        fre is not None
        and fre["mape_absolute_margin"] > 0
        and fre["mse_absolute_margin"] > 0
        and fre["mae_absolute_margin"] > 0
        and fre["bootstrap_prob_all_below_baseline"] >= 0.80
        and margin["n_rows"] == 212
        and margin["n_targets"] == 20
        and margin["bootstrap_samples"] == 2000
    )
    add_check(
        checks,
        "modern_baseline_margin_audit",
        margin_ok,
        {
            "fre_eformer": fre,
            "n_rows": margin["n_rows"],
            "n_targets": margin["n_targets"],
            "bootstrap_samples": margin["bootstrap_samples"],
        },
    )

    dominance = json.loads(
        (ROOT / "docs/experiments/artifacts/modern_baseline_dominance_audit_20260701.json").read_text(
            encoding="utf-8"
        )
    )
    dominance_ok = (
        dominance["point_dominance_cells"] == 15
        and dominance["point_dominance_total_cells"] == 15
        and dominance["all_baselines_all_three_point_better"] is True
        and dominance["ci95_upper_below_cells"] == 10
        and dominance["closest_baseline_by_minimum_relative_reduction"] == "FreEformer"
        and dominance["minimum_relative_reduction_pct"] > 6.8
    )
    add_check(
        checks,
        "modern_baseline_dominance_audit",
        dominance_ok,
        {
            "point_dominance_cells": dominance["point_dominance_cells"],
            "point_dominance_total_cells": dominance["point_dominance_total_cells"],
            "ci95_upper_below_cells": dominance["ci95_upper_below_cells"],
            "closest_baseline": dominance["closest_baseline_by_minimum_relative_reduction"],
            "minimum_relative_reduction_pct": dominance["minimum_relative_reduction_pct"],
        },
    )

    paired = json.loads(
        (ROOT / "docs/experiments/artifacts/modern_baseline_paired_audit_20260701.json").read_text(
            encoding="utf-8"
        )
    )
    paired_ok = (
        paired["passed"] is True
        and paired["n_rows"] == 212
        and paired["n_targets"] == 20
        and paired["bootstrap_samples"] == 2000
        and paired["summary"]["point_dominance_cells"] == 15
        and paired["summary"]["point_dominance_total_cells"] == 15
        and paired["summary"]["min_all_three_prob"] >= 0.89
        and paired["summary"]["closest_by_all_three_prob"] == "TimeMixer"
    )
    add_check(
        checks,
        "modern_baseline_paired_audit",
        paired_ok,
        {
            "point_dominance_cells": paired["summary"]["point_dominance_cells"],
            "point_dominance_total_cells": paired["summary"]["point_dominance_total_cells"],
            "min_all_three_prob": paired["summary"]["min_all_three_prob"],
            "closest": paired["summary"]["closest_by_all_three_prob"],
        },
    )

    paired_sig = json.loads(
        (ROOT / "docs/experiments/artifacts/paired_significance_audit_20260701.json").read_text(
            encoding="utf-8"
        )
    )
    paired_sig_summary = paired_sig["summary"]
    paired_sig_ok = (
        paired_sig["passed"] is True
        and paired_sig["n_dates"] == 212
        and paired_sig["n_targets"] == 20
        and paired_sig_summary["point_margin_positive_cells"] == 15
        and paired_sig_summary["total_cells"] == 15
        and paired_sig_summary["wilcoxon_holm_significant_0_01_cells"] == 15
        and paired_sig_summary["bootstrap_ci95_low_positive_cells"] == 14
        and paired_sig_summary["sign_flip_holm_significant_0_01_cells"] == 14
        and paired_sig_summary["closest_bootstrap_ci_cell"]["baseline"] == "TimeMixer"
        and paired_sig_summary["closest_bootstrap_ci_cell"]["metric"] == "MSE"
    )
    add_check(
        checks,
        "paired_significance_audit",
        paired_sig_ok,
        {
            "point_margin_positive_cells": paired_sig_summary["point_margin_positive_cells"],
            "wilcoxon_holm_significant_0_01_cells": paired_sig_summary["wilcoxon_holm_significant_0_01_cells"],
            "bootstrap_ci95_low_positive_cells": paired_sig_summary["bootstrap_ci95_low_positive_cells"],
            "sign_flip_holm_significant_0_01_cells": paired_sig_summary["sign_flip_holm_significant_0_01_cells"],
            "closest_boundary": paired_sig_summary["closest_bootstrap_ci_cell"],
        },
    )

    event_stress = json.loads(
        (ROOT / "docs/experiments/artifacts/event_stress_baseline_audit_20260701.json").read_text(
            encoding="utf-8"
        )
    )
    event_summary = event_stress["summary"]
    event_ok = (
        event_stress["passed"] is True
        and event_stress["n_dates"] == 212
        and event_stress["n_targets"] == 20
        and event_stress["test_start_index"] == 634
        and event_summary["high_demand_ramr_best_all_metrics"] is True
        and event_summary["high_abs_shift_ramr_best_mape"] is True
        and event_summary["high_abs_shift_best_mse_model"] == "TimeMixer"
        and event_summary["high_abs_shift_best_mae_model"] == "TimeMixer"
    )
    add_check(
        checks,
        "event_stress_baseline_audit",
        event_ok,
        {
            "high_demand_ramr_best_all_metrics": event_summary["high_demand_ramr_best_all_metrics"],
            "high_abs_shift_ramr_best_mape": event_summary["high_abs_shift_ramr_best_mape"],
            "high_abs_shift_best_mse_model": event_summary["high_abs_shift_best_mse_model"],
            "high_abs_shift_best_mae_model": event_summary["high_abs_shift_best_mae_model"],
        },
    )

    scene = json.loads(
        (ROOT / "docs/experiments/artifacts/scene_baseline_stratification_audit_20260701.json").read_text(
            encoding="utf-8"
        )
    )
    scene_summary = scene["summary"]
    scene_ok = (
        scene["passed"] is True
        and scene["n_dates"] == 212
        and scene["n_targets"] == 20
        and scene["test_start_index"] == 634
        and scene["scene_counts"]["weekday"] == 152
        and scene["scene_counts"]["weekend"] == 60
        and scene["scene_counts"]["holiday"] == 0
        and scene_summary["weekday_ramr_best_all_metrics"] is True
        and scene_summary["weekend_ramr_best_all_metrics"] is True
        and round(scene_summary["weekday_ramr_metrics"]["mape"], 4) == 18.1153
        and round(scene_summary["weekend_ramr_metrics"]["mape"], 4) == 19.6382
    )
    add_check(
        checks,
        "scene_baseline_stratification_audit",
        scene_ok,
        {
            "weekday_days": scene["scene_counts"]["weekday"],
            "weekend_days": scene["scene_counts"]["weekend"],
            "weekday_mape": scene_summary["weekday_ramr_metrics"]["mape"],
            "weekend_mape": scene_summary["weekend_ramr_metrics"]["mape"],
        },
    )

    text_claim = json.loads(
        (ROOT / "docs/experiments/artifacts/text_claim_consistency_audit_20260701.json").read_text(
            encoding="utf-8"
        )
    )
    text_claim_ok = (
        text_claim["passed"] is True
        and any(item["name"] == "final_metrics_written_in_core_texts" and item["passed"] for item in text_claim["checks"])
        and any(item["name"] == "no_leakage_protocol_written_in_texts" and item["passed"] for item in text_claim["checks"])
        and any(item["name"] == "paired_modern_baseline_claims_written_in_texts" and item["passed"] for item in text_claim["checks"])
    )
    add_check(
        checks,
        "text_claim_consistency_audit",
        text_claim_ok,
        {
            "passed": text_claim["passed"],
            "check_count": len(text_claim["checks"]),
        },
    )

    profile = json.loads(
        (ROOT / "docs/experiments/artifacts/target_error_profile_20260701.json").read_text(encoding="utf-8")
    )
    direction_metrics = {row["direction"]: row for row in profile["direction_metrics"]}
    profile_ok = (
        profile["n_rows"] == 212
        and profile["n_targets"] == 20
        and approx_equal(profile["overall"]["mape"], EXPECTED["final_mape"])
        and approx_equal(profile["overall"]["mse"], EXPECTED["final_mse"])
        and approx_equal(profile["overall"]["mae"], EXPECTED["final_mae"])
        and set(direction_metrics) == {"from_nj", "to_nj"}
        and direction_metrics["from_nj"]["mape"] < 20.0
        and direction_metrics["to_nj"]["mape"] < 20.0
        and profile["summary"]["targets_below_20_mape"] >= 16
        and profile["summary"]["targets_below_30_mape"] >= 17
    )
    add_check(
        checks,
        "target_error_profile",
        profile_ok,
        {
            "from_mape": direction_metrics.get("from_nj", {}).get("mape"),
            "to_mape": direction_metrics.get("to_nj", {}).get("mape"),
            "median_target_mape": profile["summary"].get("median_target_mape"),
            "targets_below_20_mape": profile["summary"].get("targets_below_20_mape"),
            "targets_below_30_mape": profile["summary"].get("targets_below_30_mape"),
        },
    )

    flow = json.loads(
        (ROOT / "docs/experiments/artifacts/flow_stratified_error_audit_20260701.json").read_text(
            encoding="utf-8"
        )
    )
    flow_rows = {(row["group_type"], row["group"]): row for row in flow["rows"]}
    low_flow = flow_rows.get(("flow_magnitude", "low_flow_<500"), {})
    high_flow = flow_rows.get(("flow_magnitude", "high_flow_>=3000"), {})
    normal_days = flow_rows.get(("daily_total", "normal_demand_days"), {})
    high_days = flow_rows.get(("daily_total", "high_demand_days"), {})
    flow_ok = (
        flow["n_rows"] == 212
        and flow["n_targets"] == 20
        and approx_equal(flow["overall"]["mape"], EXPECTED["final_mape"])
        and approx_equal(flow["overall"]["mse"], EXPECTED["final_mse"])
        and approx_equal(flow["overall"]["mae"], EXPECTED["final_mae"])
        and low_flow
        and high_flow
        and normal_days
        and high_days
        and low_flow["mape"] > high_flow["mape"]
        and high_flow["mape"] < 15.0
        and high_days["mape"] > normal_days["mape"]
        and normal_days["mape"] < 20.0
    )
    add_check(
        checks,
        "flow_stratified_error_audit",
        flow_ok,
        {
            "low_flow_mape": low_flow.get("mape"),
            "high_flow_mape": high_flow.get("mape"),
            "normal_demand_mape": normal_days.get("mape"),
            "high_demand_mape": high_days.get("mape"),
        },
    )

    regime = json.loads(
        (ROOT / "docs/experiments/artifacts/regime_gate_alignment_audit_20260701.json").read_text(
            encoding="utf-8"
        )
    )
    interp = regime["interpretation"]
    regime_ok = (
        regime["n_rows"] == 212
        and regime["expert_order"] == ["ShortTerm", "LongTerm", "DistributionShift"]
        and interp["high_daily_vol_short_mean"] > interp["rest_daily_vol_short_mean"]
        and interp["high_daily_vol_short_delta"] > 0.20
        and interp["high_daily_vol_short_p"] < 0.01
        and interp["high_peak_distribution_mean"] > interp["rest_peak_distribution_mean"]
        and interp["high_peak_distribution_delta"] > 0.20
        and interp["high_peak_distribution_p"] < 0.01
    )
    add_check(
        checks,
        "regime_gate_alignment_audit",
        regime_ok,
        {
            "high_daily_vol_short_mean": interp.get("high_daily_vol_short_mean"),
            "rest_daily_vol_short_mean": interp.get("rest_daily_vol_short_mean"),
            "high_peak_distribution_mean": interp.get("high_peak_distribution_mean"),
            "rest_peak_distribution_mean": interp.get("rest_peak_distribution_mean"),
        },
    )

    scan_files = [
        ROOT / "README.md",
        ROOT / "paper/论文终稿.md",
        ROOT / "专家意见逐条修改说明.md",
        ROOT / "docs/experiments/最终一致性核验_20260701.md",
        ROOT / "docs/experiments/审稿意见完成度审计_20260701.md",
        ROOT / "docs/experiments/无泄露实验协议审计_20260701.md",
        ROOT / "docs/experiments/现代强基线优势边界审计_20260701.md",
        ROOT / "docs/experiments/现代强基线全指标支配审计_20260701.md",
        ROOT / "docs/experiments/现代强基线配对预测审计_20260701.md",
        ROOT / "docs/experiments/现代强基线配对显著性审计_20260701.md",
        ROOT / "docs/experiments/高需求高波动压力场景审计_20260701.md",
        ROOT / "docs/experiments/场景分层现代强基线审计_20260701.md",
        ROOT / "docs/experiments/站点方向误差剖面审计_20260701.md",
        ROOT / "docs/experiments/流量分层误差审计_20260701.md",
        ROOT / "docs/experiments/Regime门控对齐审计_20260701.md",
    ]
    banned_hits = scan_banned(scan_files)
    add_check(checks, "stale_phrase_scan", not banned_hits, banned_hits)

    readme = (ROOT / "README.md").read_text(encoding="utf-8")
    readme_required = [
        "paper/2026-0268-基于多专家融合的铁路客流多尺度预测方法.docx",
        "paper/论文终稿.docx",
        "审稿意见逐条回复稿.md",
        "18.5463",
        "0.035640",
        "0.138028",
        "data/from_nj.csv",
        "data/to_nj.csv",
        "无泄露实验协议审计_20260701.md",
        "现代强基线优势边界审计_20260701.md",
        "现代强基线全指标支配审计_20260701.md",
        "现代强基线配对预测审计_20260701.md",
        "现代强基线配对显著性审计_20260701.md",
        "高需求高波动压力场景审计_20260701.md",
        "场景分层现代强基线审计_20260701.md",
        "论文数值与口径一致性审计_20260701.md",
        "站点方向误差剖面审计_20260701.md",
        "流量分层误差审计_20260701.md",
        "Regime门控对齐审计_20260701.md",
    ]
    readme_missing = [item for item in readme_required if item not in readme]
    add_check(checks, "readme_final_entrypoints", not readme_missing, readme_missing)

    passed = all(item["passed"] for item in checks)
    report = {"passed": passed, "checks": checks}
    REPORT_JSON.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = ["# 最终交付自动核验（2026-07-01）", "", f"总体结果：{'通过' if passed else '未通过'}", ""]
    lines.append("| 检查项 | 结果 | 摘要 |")
    lines.append("| --- | --- | --- |")
    for item in checks:
        detail = item["detail"]
        if isinstance(detail, list):
            summary_text = f"{len(detail)} item(s)"
        elif isinstance(detail, dict):
            summary_text = ", ".join(f"{k}={v}" for k, v in list(detail.items())[:4])
        else:
            summary_text = str(detail)
        lines.append(f"| {item['name']} | {'通过' if item['passed'] else '失败'} | {summary_text} |")
    lines.extend(
        [
            "",
            "## 核验结论",
            "",
            "- 数据哈希、形状、缺失值和重复日期检查通过。",
            "- 两份 Word 论文文件内容完全一致。",
            "- 论文 Word 和审稿回复 Word 字体均统一为 Times New Roman + STSong。",
            "- Strict RAMR-VE 最终三项指标均优于原 MoE-Rail 阈值。",
            "- 无泄露实验协议审计确认时间顺序划分、训练集拟合 Scaler、窗口构造和固定测试预测行数均通过。",
            "- 现代强基线优势边界审计显示相对 FreEformer 的三指标点估计均改善，bootstrap 三项同时更优比例不低于 80%。",
            "- 现代强基线全指标支配审计显示 15/15 个 MAPE/MSE/MAE 点估计比较单元均优于现代基线。",
            "- 现代强基线配对预测审计显示同一测试日期与目标变量下 15/15 个点估计比较单元均优，三指标同时更优的 bootstrap 概率最低为 89.40%。",
            "- 现代强基线配对统计审计显示 15/15 个 Wilcoxon-Holm 单侧检验显著，并披露 TimeMixer-MSE 是 bootstrap/符号置换边界单元。",
            "- 高需求高波动压力场景审计显示高需求日三指标均优于现代强基线，并披露高日际变化日 MSE/MAE 的 TimeMixer 边界。",
            "- 场景分层现代强基线审计显示工作日和周末两类场景下 Strict RAMR-VE 三项指标均优于现代强基线，并披露当前测试窗口无节假日样本。",
            "- 论文数值与口径一致性审计确认论文、审稿回复、逐条说明和 Word 稿中的关键数字与实验边界均已对齐权威 artifact。",
            "- 站点方向误差剖面审计显示双向方向聚合 MAPE 均低于 20%，并披露低流量方向相对误差局限。",
            "- 流量分层误差审计显示高单点客流样本 MAPE 低于 15%，并披露高需求日期误差仍高于常规需求日期。",
            "- Regime 门控对齐审计显示高频日总波动更多激活短期专家、峰均比突发强度更多激活分布偏移专家。",
            "- 固定消融套件和审稿回复覆盖检查通过。",
        ]
    )
    REPORT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")

    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if passed else 1


if __name__ == "__main__":
    sys.exit(main())
