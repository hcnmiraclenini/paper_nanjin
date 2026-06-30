#!/usr/bin/env python3
"""Format 北交学报 v2 定稿 docx for 铁道学报 submission."""

from __future__ import annotations

import copy
import re
import shutil
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "manuscript/基于多专家融合的高铁客流多尺度预测方法-北京交通大学学报-v2.docx"
OUT = ROOT / "drafts/铁道学报投稿定稿_修改后.docx"
FIG5 = ROOT / "assets/_paper_docx_figures/fig5_traffic_generalization.png"

# Traffic benchmark (from v2 embedded chart)
TABLE5 = [
    ("PatchTST", "0.531", "0.353"),
    ("TimeMixer", "0.484", "0.297"),
    ("iTransformer", "0.428", "0.282"),
    ("FreEformer", "0.433", "0.291"),
    ("MoE-Rail", "0.293", "0.354"),
]

ABSTRACT_CN = (
    "高精度铁路客流预测对客运组织、运力配置与智能调度具有重要意义。"
    "铁路客流序列受跨城市出行、节假日和临时扰动影响，呈现趋势、周期和残差扰动交织的多尺度非平稳特征，"
    "单一模型难以同时兼顾常规周期与非常规波动。本文提出场景条件化非平稳分解-路由机制MoE-Rail/RAMR："
    "以时序分解思想作为专家分工先验，构建短期依赖、长期依赖与分布偏移专家，"
    "并通过日历场景与窗口级非平稳状态共同条件化门控网络，实现自适应路由。"
    "为避免测试集泄露，本文采用严格时间顺序实验协议，仅在训练集拟合归一化器，"
    "在验证集确定checkpoint、集成权重与尺度校准参数，测试集仅用于最终一次评估。"
    "南京南站双向日粒度客流实验表明，Strict RAMR-VE取得18.55%的MAPE、0.035640的MSE和0.138028的MAE，"
    "三项指标均优于原MoE-Rail结果及新增现代强基线；机制分析进一步表明，场景条件与regime条件有助于提升专家路由的稳定性和可解释性。"
)

ABSTRACT_EN = (
    "High-precision railway passenger flow forecasting is important for passenger organization, "
    "capacity allocation, and intelligent dispatching. Railway passenger flow is affected by intercity "
    "travel demand, holidays, and temporary disturbances, and therefore exhibits multi-scale "
    "non-stationarity with intertwined trend, seasonal, and residual disturbance components. This paper "
    "proposes MoE-Rail/RAMR, a scenario-conditioned non-stationary decomposition and routing mechanism. "
    "The method uses time-series decomposition as an expert specialization prior, constructs short-term, "
    "long-term, and distribution-shift experts, and conditions the gating network jointly on calendar "
    "scenarios and window-level non-stationary regimes. To avoid test-set leakage, normalization is fitted "
    "only on the training set, while checkpoints, ensemble weights, and scale calibration are selected on "
    "the validation set; the test set is used only once for final evaluation. On real bidirectional daily "
    "passenger flow data of Nanjing South Railway Station, Strict RAMR-VE achieves 18.55% MAPE, 0.035640 "
    "MSE, and 0.138028 MAE, outperforming the original MoE-Rail and newly added modern baselines."
)

KEYWORDS_CN = "铁路客流；客流预测；混合专家模型；时序分解；场景感知门控；多尺度预测"
KEYWORDS_EN = (
    "railway passenger flow; passenger flow forecasting; mixture-of-experts; "
    "time series decomposition; scenario-aware gating; multi-scale forecasting"
)

FUND_TEXT = "基金项目：江苏省重点研发计划（BE2023013）"
AUTHOR_INFO = (
    "第一作者：魏涛涛（1990—），男，讲师，硕士。E-mail：weitao.wei@jsrg.com.cn\n"
    "通信作者：李兵（1978—），男，研究员，博士。E-mail：libing@ia.ac.cn"
)


def setup_matplotlib_cn():
    plt.rcParams["font.sans-serif"] = [
        "WenQuanYi Micro Hei", "SimHei", "Noto Sans CJK SC", "DejaVu Sans"
    ]
    plt.rcParams["axes.unicode_minus"] = False


def make_fig5():
    setup_matplotlib_cn()
    FIG5.parent.mkdir(exist_ok=True)
    methods = [r[0] for r in TABLE5]
    mse = [float(r[1]) for r in TABLE5]
    mae = [float(r[2]) for r in TABLE5]
    x = np.arange(len(methods))
    w = 0.36
    fig, ax = plt.subplots(figsize=(7.0, 4.0), dpi=200)
    b1 = ax.bar(x - w / 2, mse, w, label="MSE", color="#4C72B0", edgecolor="black", linewidth=0.4)
    b2 = ax.bar(x + w / 2, mae, w, label="MAE", color="#DD8452", edgecolor="black", linewidth=0.4)
    ax.set_xticks(x)
    ax.set_xticklabels(methods, rotation=15, ha="right", fontsize=9)
    ax.set_ylabel("误差", fontsize=10)
    ax.set_xlabel("方法", fontsize=10)
    ax.legend(loc="upper right", frameon=True, fontsize=9)
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for bars in (b1, b2):
        for bar in bars:
            h = bar.get_height()
            ax.text(bar.get_x() + bar.get_width() / 2, h + 0.008, f"{h:.3f}",
                    ha="center", va="bottom", fontsize=8)
    fig.tight_layout()
    fig.savefig(FIG5, bbox_inches="tight")
    plt.close(fig)


def set_run_font(run, size=10.5, bold=False, italic=False, east_asia="宋体"):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def clear_paragraph(p):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


def set_paragraph_text(p, text, size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       indent=True, space_after=0):
    clear_paragraph(p)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    p.alignment = align
    if indent and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        pf.first_line_indent = Pt(21)
    else:
        pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = copy.deepcopy(paragraph._element)
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        set_paragraph_text(new_para, text)
    return new_para


def insert_paragraph_before(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._element.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        set_paragraph_text(new_para, text, size=10.5, indent=True)
    return new_para


def insert_table_after(doc, paragraph, headers, rows):
    """Insert table immediately after paragraph (not at document end)."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
        for p in tbl.rows[0].cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.rows[ri + 1].cells[ci].text = str(val)
            for p in tbl.rows[ri + 1].cells[ci].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run)
    tbl_el = tbl._tbl
    doc.element.body.remove(tbl_el)
    paragraph._element.addnext(tbl_el)
    return tbl



def insert_paragraph_after_element(element, doc, text=""):
    """Insert empty paragraph after a block-level element (table or paragraph)."""
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    para = Paragraph(new_p, doc)
    if text:
        set_paragraph_text(para, text, size=10.5, indent=True)
    return para


def replace_inline_image(paragraph, img_path, width_cm=14.0):
    """Replace first inline image in paragraph."""
    drawings = paragraph._element.findall(".//" + qn("w:drawing"))
    if not drawings:
        run = paragraph.add_run()
        run.add_picture(str(img_path), width=Cm(width_cm))
        return
    # remove old drawing by clearing runs with blip
    for r in list(paragraph.runs):
        if "graphic" in r._element.xml or "drawing" in r._element.xml:
            r._element.getparent().remove(r._element)
    run = paragraph.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))


def apply_text_fixes(text: str) -> str:
    fixes = [
        ("\\尽管如此", "尽管如此"),
        ("表格1", "表1"),
        ("表格 1", "表1"),
        ("MoE-rail", "MoE-Rail"),
        ("MoE-rail", "MoE-Rail"),
        ("T=7天", "T=6天"),
        ("T=7 天", "T=6 天"),
        ("历史窗口T=7", "历史窗口T=6"),
        ("7个具有代表性的方法", "6个具有代表性的对比基线"),
        ("选取了如下的7个", "选取了如下的6个"),
        ("如图2所示。MoE-Rail在traffic", "如图5所示。MoE-Rail在Traffic"),
        ("表3报告了上述四个预测窗口", "表5报告了上述四个预测窗口"),
        ("表3 报告了上述四个预测窗口", "表5报告了上述四个预测窗口"),
        ("MAPE降至约19%", "MAPE降至19.41%"),
        ("MAPE 降至约 19%", "MAPE降至19.41%"),
        ("动态门控网络", "场景感知门控"),
        ("动态门控", "场景感知门控"),
        ("长期周期专家", "长期依赖专家"),
        ("Fig3:", "Fig. 5 "),
        ("Fig. 5  Distribution of expert gating weights", "Fig. 3  Distribution of expert gating weights"),
        ("表3 ：", "表5 "),
        ("图3 ：MoE-Rail在traffic", "图5  MoE-Rail在Traffic"),
        ("图3：MoE-Rail在traffic", "图5  MoE-Rail在Traffic"),
        ("（2.7都是新加入的）", ""),
        ("（确认无误把上面这一段摘要删除）", ""),
    ]
    for old, new in fixes:
        text = text.replace(old, new)
    # normalize figure refs in 2.8 only - careful with 图3 in 2.7
    return text


def format_heading(p, level=2):
    text = p.text.strip()
    if not text:
        return
    sizes = {1: 14, 2: 12, 3: 11}
    sz = sizes.get(level, 12)
    set_paragraph_text(p, text, size=sz, bold=True, indent=False, space_after=3)
    p.paragraph_format.space_before = Pt(6)


def detect_heading_level(text):
    t = text.strip()
    if re.match(r"^(\d+\s+|\d+\.\d+\s)", t):
        if re.match(r"^\d+\.\d+\.\d+", t):
            return 3
        if re.match(r"^\d+\.\d+", t):
            return 3
        if re.match(r"^2\.\d+", t) or re.match(r"^1\.\d+", t):
            return 3 if t.count(".") >= 2 else 2
        if t.startswith(("0 ", "1 ", "2 ", "3 ")):
            return 2
    if t in ("引言", "结论", "参考文献", "预测模型的建立"):
        return 2
    return 0


def is_table_caption(text):
    return bool(re.match(r"^表\s*\d+", text.strip()))


def is_figure_caption(text):
    return bool(re.match(r"^图\s*\d+", text.strip()))


def is_english_caption(text):
    t = text.strip()
    return t.startswith(("Table ", "Fig.", "Fig "))


FIG5_NOTE = (
    "图5进一步展示了各方法在Traffic数据集上的MSE/MAE对比，"
    "MoE-Rail在MSE维度保持优势，但MAE并未全面优于iTransformer、TimeMixer和FreEformer；"
    "因此该结果仅用于说明重构小时级窗口和周期特征后的结构鲁棒性，不能作为稳定跨域泛化或所有指标全面占优的证据。"
)


def is_fig5_traffic_caption(text: str) -> bool:
    t = text.strip()
    return bool(re.match(r"^图5\s+MoE-Rail.*Traffic", t))


def cleanup_section_28(doc):
    """Remove duplicate Traffic captions and duplicate §2.8 analysis lines."""
    dup_indices = []
    seen_fig5 = False
    seen_t5_ana = False
    seen_fig5_note = False
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if "图5进一步展示" in t:
            if seen_fig5_note:
                dup_indices.append(i)
            else:
                seen_fig5_note = True
        if t.startswith("由表5可见"):
            if seen_t5_ana:
                dup_indices.append(i)
            else:
                seen_t5_ana = True
        if is_fig5_traffic_caption(t):
            if seen_fig5:
                dup_indices.append(i)
                if i + 1 < len(doc.paragraphs) and doc.paragraphs[i + 1].text.strip().startswith("Fig"):
                    dup_indices.append(i + 1)
            else:
                seen_fig5 = True
    for i in sorted(set(dup_indices), reverse=True):
        doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

    # Remove drawings embedded inside text paragraphs (keep dedicated image paragraph only)
    for p in doc.paragraphs:
        if p._element.findall(".//" + qn("w:drawing")) and len(p.text.strip()) > 20:
            for d in p._element.findall(".//" + qn("w:drawing")):
                d.getparent().remove(d)


def find_fig5_traffic_index(doc):
    for i, p in enumerate(doc.paragraphs):
        if is_fig5_traffic_caption(p.text):
            return i
    return None


def normalize_captions(doc):
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        m = re.match(r"^(表\d+):", t)
        if m:
            new_t = t.replace(":", "  ", 1)
            set_paragraph_text(p, new_t, size=10.5, bold=True,
                               align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        elif re.match(r"^Table\s+\d+:", t):
            new_t = t.replace(":", "  ", 1)
            set_paragraph_text(p, new_t, size=10.5,
                               align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        elif re.match(r"^Fig\.?\s*\d", t, re.I):
            num_m = re.match(r"^Fig\.?\s*(\d+)", t, re.I)
            num = int(num_m.group(1)) if num_m else 0
            if num == 3:
                new_t = "Fig. 3  Distribution of expert gating weights under calendar scenarios"
            elif num == 4:
                new_t = "Fig. 4  Temporal alignment between STL components and expert gating weights"
            elif num == 5:
                new_t = "Fig. 5  Performance of MoE-Rail on Traffic dataset (four-horizon average)"
            else:
                new_t = re.sub(r"^Fig\.?\s*(\d+)", r"Fig. \1", t, flags=re.I)
            set_paragraph_text(p, new_t, size=10.5,
                               align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)


def finalize_section_28(doc):
    cleanup_section_28(doc)
    normalize_captions(doc)

    fig5_idx = find_fig5_traffic_index(doc)
    if fig5_idx is None:
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith("3") and "结论" in p.text:
                intro = doc.paragraphs[i - 1] if i > 0 else p
                if not any("图5进一步展示" in x.text for x in doc.paragraphs):
                    note = insert_paragraph_before(p, FIG5_NOTE)
                    set_paragraph_text(note, FIG5_NOTE, size=10.5, indent=True)
                    intro = note
                img_p = insert_paragraph_before(p, "")
                replace_inline_image(img_p, FIG5, width_cm=14.0)
                cap = insert_paragraph_before(p, "图5  MoE-Rail在Traffic数据集上的表现（四步长平均）")
                set_paragraph_text(cap, cap.text, size=10.5, bold=True,
                                   align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
                cap_en = insert_paragraph_before(p,
                    "Fig. 5  Performance of MoE-Rail on Traffic dataset (four-horizon average)")
                set_paragraph_text(cap_en, cap_en.text, size=10.5,
                                   align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
                break
        return

    # Reorder: analysis note should precede image, image immediately precedes caption
    fig5_idx = find_fig5_traffic_index(doc)
    cap = doc.paragraphs[fig5_idx]
    note_idx = next((i for i, p in enumerate(doc.paragraphs) if "图5进一步展示" in p.text), None)
    if note_idx is not None and note_idx > fig5_idx:
        note_el = doc.paragraphs[note_idx]._element
        cap._element.addprevious(note_el)

    # Remove orphaned image paragraphs after Fig. 5 English caption
    rm = []
    for i in range(fig5_idx + 2, min(fig5_idx + 6, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if t.startswith("3") and "结论" in t:
            break
        if p._element.findall(".//" + qn("w:drawing")) and not t:
            rm.append(i)
    for i in sorted(rm, reverse=True):
        doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

    fig5_idx = find_fig5_traffic_index(doc)
    cap = doc.paragraphs[fig5_idx]

    if not any("图5进一步展示" in p.text for p in doc.paragraphs):
        note = insert_paragraph_before(cap, FIG5_NOTE)
        set_paragraph_text(note, FIG5_NOTE, size=10.5, indent=True)
        fig5_idx = find_fig5_traffic_index(doc)
        cap = doc.paragraphs[fig5_idx]

    # Keep exactly one image paragraph immediately before 图5 Traffic caption
    fig5_idx = find_fig5_traffic_index(doc)
    cap = doc.paragraphs[fig5_idx]
    sec_start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("2.8"):
            sec_start = i
            break
    img_before = []
    if sec_start is not None:
        for i in range(sec_start, fig5_idx):
            if doc.paragraphs[i]._element.findall(".//" + qn("w:drawing")):
                img_before.append(i)

    if img_before:
        keep = img_before[-1]
        for i in img_before[:-1]:
            el = doc.paragraphs[i]._element
            for d in el.findall(".//" + qn("w:drawing")):
                d.getparent().remove(d)
            if not doc.paragraphs[i].text.strip():
                el.getparent().remove(el)
        fig5_idx = find_fig5_traffic_index(doc)
        if fig5_idx - 1 != keep:
            # move image to paragraph right before caption
            prev = doc.paragraphs[fig5_idx - 1]
            if not prev._element.findall(".//" + qn("w:drawing")):
                if prev.text.strip():
                    img_p = insert_paragraph_before(doc.paragraphs[fig5_idx], "")
                else:
                    img_p = prev
                replace_inline_image(img_p, FIG5, width_cm=14.0)
            else:
                replace_inline_image(prev, FIG5, width_cm=14.0)
        else:
            replace_inline_image(doc.paragraphs[keep], FIG5, width_cm=14.0)
    else:
        prev = doc.paragraphs[fig5_idx - 1] if fig5_idx > 0 else cap
        if prev.text.strip():
            img_p = insert_paragraph_before(cap, "")
        else:
            img_p = prev
        replace_inline_image(img_p, FIG5, width_cm=14.0)

    cleanup_section_28(doc)


def process_document():
    make_fig5()
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)

    # --- Pass 1: delete unwanted paragraphs (iterate backwards) ---
    to_delete = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        if "确认无误" in t:
            to_delete.append(i)
        if t == "（2.7都是新加入的）":
            to_delete.append(i)
        # old duplicate abstract (v2 opening)
        if t.startswith("摘  要：") and "优化客运组织" in t:
            to_delete.append(i)
        if t.startswith("摘要：") and "优化客运组织" in t:
            to_delete.append(i)
        # misplaced keyword line under 中图分类号
        if t.startswith("铁路客流；") and "由编辑部填写" in t:
            to_delete.append(i)
        if t == "中图分类号":
            to_delete.append(i)

    for i in sorted(set(to_delete), reverse=True):
        el = doc.paragraphs[i]._element
        el.getparent().remove(el)

    # --- Title / authors / affiliation ---
    paras = doc.paragraphs
    if paras[0].text.strip():
        set_paragraph_text(paras[0], paras[0].text.strip(), size=16, bold=True,
                           align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    for idx in [2, 3]:
        if idx < len(paras) and paras[idx].text.strip():
            set_paragraph_text(paras[idx], paras[idx].text.strip(), size=16,
                               align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    if len(paras) > 4 and paras[4].text.strip():
        set_paragraph_text(paras[4], paras[4].text.strip(), size=9,
                           align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    # --- Pass 2: front matter fixes ---
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        if i < 15 and (
            t.startswith("摘要：")
            or (t.startswith("铁路已成为") and "MoE-Rail" in t)
            or (t.startswith("高精度铁路客流") and "MoE-Rail" in t)
        ):
            set_paragraph_text(p, f"摘要：{ABSTRACT_CN}", size=9, indent=False)
        elif t.startswith("关键词"):
            set_paragraph_text(p, f"关键词：{KEYWORDS_CN}", size=9, indent=False)
        elif t.startswith("DOI"):
            set_paragraph_text(p, "DOI：（由编辑部填写）", size=9, indent=False)
        elif t.startswith("Abstract"):
            set_paragraph_text(p, f"Abstract: {ABSTRACT_EN}", size=10.5, indent=False)
        elif t.startswith("Keywords"):
            set_paragraph_text(p, f"Keywords: {KEYWORDS_EN}", size=10.5, indent=False)

    # Insert 中图分类号 if missing
    inserted_cls = any(p.text.strip().startswith("中图分类号：") for p in doc.paragraphs[:15])
    if not inserted_cls:
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith("关键词"):
                cp = insert_paragraph_after(p, "中图分类号：U293.1    文献标志码：A")
                set_paragraph_text(cp, cp.text, size=9, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
                break
    else:
        for p in doc.paragraphs[:15]:
            if p.text.strip().startswith("中图分类号"):
                set_paragraph_text(p, "中图分类号：U293.1    文献标志码：A", size=9,
                                   indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Remove duplicate metadata lines in front matter
    seen_meta = set()
    meta_delete = []
    for i, p in enumerate(doc.paragraphs[:20]):
        t = p.text.strip()
        key = None
        if t.startswith("中图分类号"):
            key = "cls"
        elif t.startswith("DOI"):
            key = "doi"
        elif t.startswith("关键词"):
            key = "kw"
        if key:
            if key in seen_meta:
                meta_delete.append(i)
            else:
                seen_meta.add(key)
    for i in sorted(meta_delete, reverse=True):
        doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

    # --- Pass 3: body text fixes + formatting ---
    in_front = True
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        if t == "引言":
            in_front = False

        if in_front:
            continue

        # apply text fixes
        new_t = apply_text_fixes(t)
        if new_t != t:
            drawings = p._element.findall(".//" + qn("w:drawing"))
            if not drawings:
                hl = detect_heading_level(new_t)
                if hl:
                    format_heading(p, hl)
                elif is_table_caption(new_t) or is_figure_caption(new_t) or is_english_caption(new_t):
                    set_paragraph_text(p, new_t, size=10.5, bold=True,
                                       align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
                else:
                    set_paragraph_text(p, new_t, size=10.5, indent=True)

        t = p.text.strip()
        hl = detect_heading_level(t)
        if hl and not p._element.findall(".//" + qn("w:drawing")):
            format_heading(p, hl)
        elif is_table_caption(t) or is_figure_caption(t):
            set_paragraph_text(p, re.sub(r"表\s*", "表", t).replace("：", "  "), size=10.5,
                               bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        elif is_english_caption(t):
            set_paragraph_text(p, t, size=10.5, bold=False,
                               align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        elif t.startswith("参考文献"):
            format_heading(p, 2)
        elif not p._element.findall(".//" + qn("w:drawing")) and hl == 0:
            if not (is_table_caption(t) or is_figure_caption(t) or is_english_caption(t)):
                # body paragraph
                pf = p.paragraph_format
                pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                pf.first_line_indent = Pt(21)
                for run in p.runs:
                    set_run_font(run, size=10.5)

    # --- Insert 2.7.2 heading if missing ---
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("2.7.1"):
            nxt = doc.paragraphs[i + 1].text.strip() if i + 1 < len(doc.paragraphs) else ""
            if not nxt.startswith("2.7.2"):
                anchor = doc.paragraphs[i + 1]
                hp = insert_paragraph_after(anchor, "2.7.2  日历场景下的门控权重统计")
                format_heading(hp, 3)
            break

    # --- Fix 2.8: insert table5 block if needed ---
    has_table5_caption = any(p.text.strip().startswith("表5  Traffic") for p in doc.paragraphs)
    has_t5_analysis = any("由表5可见" in p.text for p in doc.paragraphs)
    if not has_table5_caption:
        for i, p in enumerate(doc.paragraphs):
            t = p.text.strip()
            if is_fig5_traffic_caption(t):
                anchor = doc.paragraphs[i - 2] if i >= 2 else p
                t5_intro = insert_paragraph_after(
                    anchor,
                    "表5汇总四个预测步长下各方法测试结果的平均值，用以评估MoE-Rail在重构小时级窗口和周期特征后的结构鲁棒性。",
                )
                set_paragraph_text(t5_intro, t5_intro.text, size=10.5, indent=True)
                cap_t5 = insert_paragraph_after(t5_intro, "表5  Traffic数据集四个预测步长平均测试结果")
                set_paragraph_text(cap_t5, cap_t5.text, size=10.5, bold=True,
                                   align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
                cap_t5_en = insert_paragraph_after(
                    cap_t5, "Table 5  Average Traffic benchmark results across four horizons")
                set_paragraph_text(cap_t5_en, cap_t5_en.text, size=10.5,
                                   align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
                tbl5 = insert_table_after(doc, cap_t5_en, ["方法", "MSE", "MAE"], TABLE5)
                t5_ana = insert_paragraph_after_element(tbl5._tbl, doc,
                    "由表5可见，MoE-Rail平均MSE为0.293，为各对比方法中最低，较PatchTST（0.531）"
                    "降低44.8%；平均MAE为0.354，尚未达到对比方法中的最优水平。实验结果如图5所示。")
                set_paragraph_text(t5_ana, t5_ana.text, size=10.5, indent=True)
                fig_note = insert_paragraph_after(t5_ana, FIG5_NOTE)
                set_paragraph_text(fig_note, FIG5_NOTE, size=10.5, indent=True)
                img_p = insert_paragraph_after(fig_note, "")
                replace_inline_image(img_p, FIG5, width_cm=14.0)
                break
    elif not has_t5_analysis:
        tbl_el = None
        for p in doc.paragraphs:
            if p.text.strip().startswith("Table 5"):
                el = p._element.getnext()
                while el is not None:
                    if el.tag.endswith("tbl"):
                        tbl_el = el
                        break
                    el = el.getnext()
                break
        if tbl_el is not None:
            t5_ana = insert_paragraph_after_element(tbl_el, doc,
                "由表5可见，MoE-Rail平均MSE为0.293，为各对比方法中最低，较PatchTST（0.531）"
                "降低44.8%；平均MAE为0.354，尚未达到对比方法中的最优水平。实验结果如图5所示。")
            set_paragraph_text(t5_ana, t5_ana.text, size=10.5, indent=True)
            fig_note = insert_paragraph_after(t5_ana, FIG5_NOTE)
            set_paragraph_text(fig_note, fig_note.text, size=10.5, indent=True)
            img_p = insert_paragraph_after(fig_note, "")
            replace_inline_image(img_p, FIG5, width_cm=14.0)

    # Fix mis-labeled Fig.3 caption if corrupted
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Fig. 5") and "expert gating" in t.lower():
            set_paragraph_text(p, "Fig. 3  Distribution of expert gating weights under calendar scenarios",
                               size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if is_fig5_traffic_caption(t):
            set_paragraph_text(p, "图5  MoE-Rail在Traffic数据集上的表现（四步长平均）",
                               size=10.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
            if i + 1 < len(doc.paragraphs):
                set_paragraph_text(
                    doc.paragraphs[i + 1],
                    "Fig. 5  Performance of MoE-Rail on Traffic dataset (four-horizon average)",
                    size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
            break

    # --- Fix 2.2 baseline paragraph ---
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("为了全面验证") and "7个" in t:
            new = apply_text_fixes(t)
            new = new.replace(
                "我们选取了如下的6个具有代表性的对比基线作为对比方法，包括",
                "选取6个具有代表性的对比基线，包括"
            )
            set_paragraph_text(p, new, size=10.5, indent=True)

    # --- Add fund + author info before 引言 ---
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "引言":
            fp = insert_paragraph_after(doc.paragraphs[i - 1], FUND_TEXT)
            set_paragraph_text(fp, FUND_TEXT, size=7.5, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
            ap = insert_paragraph_after(fp, AUTHOR_INFO)
            set_paragraph_text(ap, AUTHOR_INFO, size=7.5, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
            break

    # --- Enhance conclusion section heading ---
    for p in doc.paragraphs:
        t = p.text.strip()
        if re.match(r"^3\s+结论", t):
            format_heading(p, 2)

    # Update 2.5 opening if still says wrong
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("图2对") and "表格" in t:
            set_paragraph_text(p, apply_text_fixes(t), size=10.5, indent=True)

    # Fix reference [2] missing number for 赵阳阳
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("赵阳阳") and not t.startswith("[2]"):
            set_paragraph_text(p, "[2] " + t, size=10.5, indent=False)

    finalize_section_28(doc)

    doc.save(OUT)
    return OUT


def verify(out: Path):
    d = Document(out)
    text = "\n".join(p.text for p in d.paragraphs)
    issues = []
    if "图3 ：MoE-Rail在traffic" in text or "图3：MoE-Rail在traffic" in text:
        issues.append("Traffic figure still labeled 图3")
    if "表3报告" in text and "2.8" in text:
        issues.append("2.8 still references 表3")
    if "T=7" in text:
        issues.append("T=7 not fixed")
    if "7个具有" in text:
        issues.append("7 baselines not fixed")
    if len(d.tables) < 5:
        issues.append(f"expected 5 tables, got {len(d.tables)}")
    refs = re.findall(r"图\s*\d+", text)
    print("Figure refs:", sorted(set(refs)))
    print("Tables:", len(d.tables), "Images:", sum(1 for r in d.part.rels.values() if "image" in r.reltype))
    if issues:
        print("ISSUES:", issues)
    else:
        print("Verification passed.")
    return issues


if __name__ == "__main__":
    out = process_document()
    print(f"Saved: {out}")
    verify(out)
